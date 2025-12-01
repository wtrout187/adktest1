"""
Enterprise CLM System with ChromaDB, RAG, and PDF Support
=========================================================

Enhanced version with:
- ChromaDB vector database for persistent storage
- PDF parsing and upload capabilities
- RAG (Retrieval-Augmented Generation) for context-aware analysis
- Parallel processing support for bulk contract loading
- Optimized for 16GB RAM CPU-only systems

Author: ADKTest Project
"""

import os
import re
import asyncio
import json
import hashlib
import shutil
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime

# Load environment variables
from dotenv import load_dotenv
load_dotenv(Path(__file__).parent / "my_agent" / ".env")


def _load_env_list(var_name: str, lowercase: bool = False) -> List[str]:
    """Return a cleaned list from a pipe-delimited environment variable."""
    raw_value = os.getenv(var_name, "")
    values = [item.strip() for item in raw_value.split('|') if item.strip()]
    if lowercase:
        values = [value.lower() for value in values]
    return values


def _clean_env(value: Optional[str], fallback: str) -> str:
    """Trim whitespace and provide a fallback when env vars are blank."""
    if value is None:
        return fallback
    stripped = value.strip()
    return stripped or fallback


PRIMARY_PARTY_PREFIX = os.getenv('CLM_PRIMARY_PARTY_PREFIX', 'primary_party').strip().lower() or 'primary_party'
PRIMARY_COMPANY_NAMES = set(_load_env_list('CLM_PRIMARY_COMPANY_NAMES', lowercase=True))
PRIMARY_COMPANY_ADDRESSES = _load_env_list('CLM_PRIMARY_COMPANY_ADDRESSES', lowercase=True)
PRIMARY_COMPANY_ROLE_HINTS = set(_load_env_list('CLM_PRIMARY_COMPANY_ROLE_HINTS', lowercase=True))
PRIMARY_COMPANY_EMEA_PATTERN = os.getenv('CLM_PRIMARY_COMPANY_EMEA_ENTITY_PATTERN', '')
PRIMARY_COMPANY_GDPR_REASON = os.getenv('CLM_PRIMARY_COMPANY_GDPR_REASON', 'Primary EMEA entity present')

OWNER_LEGAL = os.getenv('CLM_OWNER_LEGAL', '')
OWNER_DISTRIBUTION_NA = os.getenv('CLM_OWNER_DISTRIBUTION_NA', '')
OWNER_DISTRIBUTION_EMEA = os.getenv('CLM_OWNER_DISTRIBUTION_EMEA', '')
OWNER_DISTRIBUTION_NA_APAC = os.getenv('CLM_OWNER_DISTRIBUTION_NA_APAC', 'tom.roberts@spectralink.com')
OWNER_MARKETING = os.getenv('CLM_OWNER_MARKETING', '')
OWNER_PROCUREMENT = os.getenv('CLM_OWNER_PROCUREMENT', '')

DEFAULT_CONTRACT_OWNER = os.getenv('CLM_DEFAULT_CONTRACT_OWNER', OWNER_LEGAL or 'wayne.trout@spectralink.com')
DEFAULT_CONTRACT_CO_OWNER = os.getenv('CLM_DEFAULT_CONTRACT_CO_OWNER', 'wayne.trout@spectralink.com')

MODEL_PROVIDER = _clean_env(os.getenv('CLM_MODEL_PROVIDER'), 'local').lower()
LOCAL_MODEL_NAME_RAW = _clean_env(os.getenv('CLM_LOCAL_MODEL_NAME'), '')
LOCAL_MODEL_NAME = LOCAL_MODEL_NAME_RAW or None
GEMINI_MODEL_NAME = _clean_env(os.getenv('CLM_GEMINI_MODEL'), 'gemini-2.5-flash')
EMBEDDING_MODEL_OVERRIDE = _clean_env(os.getenv('CLM_EMBEDDING_MODEL'), '') or None

if not PRIMARY_COMPANY_NAMES:
    print("⚠️  CLM_PRIMARY_COMPANY_NAMES is not set; primary party detection may be limited.")
if not PRIMARY_COMPANY_ADDRESSES:
    print("⚠️  CLM_PRIMARY_COMPANY_ADDRESSES is not set; address-based detection may be limited.")
if not OWNER_LEGAL:
    print("⚠️  CLM_OWNER_LEGAL is not set; SharePoint owner routing will omit primary owner.")

STANDARD_METADATA_DEFAULTS = {
    'contract_type': 'Auto-detected',
    'parties': 'Unknown',
    'gdpr_applicable': 'Unknown',
    'gdpr_reason': '',
    'region': 'Unknown',
    'sharepoint_department': 'Legal' if OWNER_LEGAL else 'General',
    'sharepoint_owner_primary': OWNER_LEGAL,
    'termination_for_convenience': 'Unknown',
    'governing_law': 'Unknown',
    'effective_date': 'Unknown',
    'termination_date': 'Unknown',
    'initial_term': 'Unknown',
    'renewal_term': 'Unknown'
}

BOOLEAN_METADATA_DEFAULTS: Dict[str, bool] = {
    'auto_renewal': False,
    'perpetual': False,
    'has_related_documents': False,
    'has_tables': False,
    'ocr_applied': False,
    'is_draft': False,
    'is_archive': False,
    'is_terminated': False,
    'subject_to_gdpr': False
}


def _filter_owners(*owners: str) -> List[str]:
    filtered: List[str] = []
    for owner in owners:
        if owner and owner not in filtered:
            filtered.append(owner)
    return filtered


def _ensure_owner_defaults(owner_list: List[str]) -> List[str]:
    owners = [owner for owner in owner_list if owner]
    if not owners and DEFAULT_CONTRACT_OWNER:
        owners.append(DEFAULT_CONTRACT_OWNER)
    if DEFAULT_CONTRACT_CO_OWNER and DEFAULT_CONTRACT_CO_OWNER not in owners:
        owners.append(DEFAULT_CONTRACT_CO_OWNER)
    return owners


def _ensure_boolean_metadata_defaults(metadata: Dict[str, Any]) -> None:
    """Ensure key boolean metadata fields are always populated."""
    for field, default in BOOLEAN_METADATA_DEFAULTS.items():
        if field not in metadata:
            metadata[field] = default


def _add_primary_party_field(details: Dict[str, str], field: str, value: Optional[str]) -> None:
    if not value:
        return
    base_key = f'primary_party_{field}'
    details[base_key] = value
    if PRIMARY_PARTY_PREFIX != 'primary_party':
        details[f'{PRIMARY_PARTY_PREFIX}_{field}'] = value


def _humanize_key(key: str) -> str:
    parts = key.split('_')
    humanized_parts = []
    for part in parts:
        lower = part.lower()
        if lower == 'gdpr':
            humanized_parts.append('GDPR')
        elif lower == 'msa':
            humanized_parts.append('MSA')
        else:
            humanized_parts.append(lower.capitalize())
    return ' '.join(humanized_parts)


def _sanitize_metadata_value(value: Any) -> Optional[Any]:
    if value is None:
        return None
    if isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, (list, tuple, set)):
        flattened = [str(item).strip() for item in value if item not in (None, '', [])]
        return ', '.join(flattened) if flattened else None
    if isinstance(value, dict):
        try:
            return json.dumps(value, ensure_ascii=False)
        except Exception:
            return str(value)
    return str(value)


def _sanitize_metadata_dict(metadata: Dict[str, Any]) -> Dict[str, Any]:
    sanitized: Dict[str, Any] = {}
    for key, value in metadata.items():
        cleaned_value = _sanitize_metadata_value(value)
        if cleaned_value is None:
            continue
        sanitized[str(key)] = cleaned_value
    return sanitized


def _slugify_filename(name: str) -> str:
    slug = re.sub(r'[^a-z0-9]+', '-', name.lower()).strip('-')
    return slug or 'contract'


def _compute_source_hash(file_path: str, text: str) -> str:
    hasher = hashlib.sha1()
    try:
        path_obj = Path(file_path)
        hasher.update(path_obj.name.lower().encode('utf-8'))
        stats = path_obj.stat()
        hasher.update(str(int(stats.st_size)).encode('utf-8'))
        hasher.update(str(int(stats.st_mtime)).encode('utf-8'))
        with path_obj.open('rb') as fh:
            chunk = fh.read(262144)
            hasher.update(chunk)
    except Exception:
        hasher.update(Path(file_path).name.lower().encode('utf-8'))
    sample = text[:2000].encode('utf-8', 'ignore')
    hasher.update(sample)
    return hasher.hexdigest()


def _generate_contract_id(file_path: str, text: str, source_hash: Optional[str] = None) -> str:
    hash_value = source_hash or _compute_source_hash(file_path, text)
    slug = _slugify_filename(Path(file_path).stem)[:40]
    return f"{slug}-{hash_value[:10]}"


def _extract_contract_id_from_message(message: str) -> Optional[str]:
    match = re.search(r"Contract ID:\s*([\w\-]+)", message)
    return match.group(1) if match else None


def _ensure_dir_on_path(directory: Path) -> None:
    try:
        dir_str = str(directory)
        current_path = os.environ.get('PATH', '')
        segments = [segment.strip() for segment in current_path.split(os.pathsep) if segment]
        if dir_str not in segments:
            os.environ['PATH'] = os.pathsep.join([dir_str] + segments)
    except Exception as path_err:
        print(f"⚠️  Could not append {directory} to PATH: {path_err}")


DATE_LITERAL_PATTERNS = [
    # English month names (full and abbreviated)
    r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{1,2},?\s+\d{4}',
    # Day-Month-Year with month name
    r'\b\d{1,2}\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}\b',
    # Numeric formats
    r'\b\d{1,2}/\d{1,2}/\d{2,4}\b',
    r'\b\d{1,2}-\d{1,2}-\d{2,4}\b',
    r'\b\d{1,2}\.\d{1,2}\.\d{2,4}\b',
    r'\b\d{4}-\d{2}-\d{2}\b',
    # Year only in filename context (e.g., "Oct 19, 2005" from folder name)
    r'\b\d{4}\b',
]


def _extract_date_literal(text: str) -> Optional[str]:
    for pattern in DATE_LITERAL_PATTERNS:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            extracted = match.group(0).replace('  ', ' ').strip(' ,.;')
            
            # Validate: if it's a 4-digit number, ensure it's a reasonable year (1900-2099)
            if re.match(r'^\d{4}$', extracted):
                year = int(extracted)
                if year < 1900 or year > 2099:
                    continue  # Skip invalid years, try next pattern
            
            return extracted
    return None


def _format_numeric_token_as_date(token: str) -> Optional[str]:
    token = re.sub(r'[^0-9]', '', token)
    if not token:
        return None
    candidates = []
    if len(token) == 6:
        candidates = ["%m%d%y"]
    elif len(token) == 8:
        candidates = ["%Y%m%d", "%m%d%Y"]
    elif len(token) == 4:
        # treat yyyy
        try:
            year = int(token)
            if 1900 <= year <= 2100:
                return f"{year}-01-01"
        except ValueError:
            return None
    else:
        return None

    for fmt in candidates:
        try:
            parsed = datetime.strptime(token, fmt)
            if 1900 <= parsed.year <= 2100:
                return parsed.strftime("%Y-%m-%d")
        except ValueError:
            continue
    return None

# ADK imports (optional)
try:
    from google.adk.agents import Agent
    from google.adk.models import Gemini
    from google.adk.runners import Runner
    from google.adk.sessions import InMemorySessionService
    from google.adk.tools.tool_context import ToolContext
    from google.adk.agents.callback_context import CallbackContext
    from google.adk.models.llm_request import LlmRequest
    from google.adk.models.llm_response import LlmResponse
    from google.adk.tools.base_tool import BaseTool
    from google.genai.types import Content, Part  # For proper message formatting
    ADK_AVAILABLE = True
    try:
        from local_llm_adapter import LocalOllamaLlm
        LOCAL_LLM_AVAILABLE = True
    except Exception as local_adapter_err:  # noqa: BLE001 - surface config issues
        LOCAL_LLM_AVAILABLE = False
        print(f"⚠️  Local model adapter unavailable: {local_adapter_err}")
except ImportError:
    Agent = Gemini = Runner = InMemorySessionService = ToolContext = Any  # type: ignore
    CallbackContext = LlmRequest = LlmResponse = BaseTool = Any  # type: ignore
    Content = Part = Any  # type: ignore
    LocalOllamaLlm = Any  # type: ignore
    ADK_AVAILABLE = False
    LOCAL_LLM_AVAILABLE = False
    print("[!] google-adk not installed. Agent/chat features will be disabled.")

# Vector database & embeddings
try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    print("⚠️  ChromaDB not installed. Install with: pip install chromadb")

# PDF parsing with OCR support
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠️  PyMuPDF not installed. Install with: pip install pymupdf")

# LLM-optimized PDF extraction
try:
    import pymupdf4llm
    PYMUPDF4LLM_AVAILABLE = True
except ImportError:
    PYMUPDF4LLM_AVAILABLE = False

try:
    import ocrmypdf
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False
    print("⚠️  ocrmypdf not installed. Install with: pip install ocrmypdf")

# Word document support
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("⚠️  python-docx not installed. Install with: pip install python-docx")

# Embeddings
try:
    from sentence_transformers import SentenceTransformer
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    EMBEDDINGS_AVAILABLE = False
    print("⚠️  sentence-transformers not installed. Install with: pip install sentence-transformers")

# Optional fuzzy matching accelerator for party detection
try:
    from rapidfuzz import fuzz, process
    RAPIDFUZZ_AVAILABLE = True
except ImportError:
    RAPIDFUZZ_AVAILABLE = False
    print("⚠️  rapidfuzz not installed. Install with: pip install rapidfuzz")

DEFAULT_TESSERACT_DIR = os.getenv('CLM_TESSERACT_PATH', r"C:\Users\wtrout\Python Projects\alicecms2\tesseract")
DEFAULT_GHOSTSCRIPT_HINT = os.getenv('CLM_GHOSTSCRIPT_PATH', r"C:\Program Files\gs")


def _configure_tesseract() -> Optional[str]:
    def _normalize_path(value: str) -> Path:
        return Path(value.strip('"')).expanduser()

    def _record_tesseract_path(executable: Path) -> str:
        os.environ['TESSERACT_CMD'] = str(executable)
        os.environ['OCRMYPDF_TESSERACT'] = str(executable)
        _ensure_dir_on_path(executable.parent)
        print(f"✅ Tesseract configured: {executable}")
        return str(executable)

    env_cmd = os.getenv('TESSERACT_CMD')
    if env_cmd:
        env_candidate = _normalize_path(env_cmd)
        if env_candidate.exists():
            return _record_tesseract_path(env_candidate)
        print(f"⚠️  TESSERACT_CMD was set but not found at {env_candidate}")

    candidate_inputs = [
        os.getenv('CLM_TESSERACT_PATH') or '',
        DEFAULT_TESSERACT_DIR,
        r"C:\\Program Files\\Tesseract-OCR",
        r"C:\\Program Files (x86)\\Tesseract-OCR"
    ]

    for candidate in candidate_inputs:
        if not candidate:
            continue
        path_candidate = _normalize_path(candidate)
        if path_candidate.is_file() and path_candidate.name.lower().startswith('tesseract'):
            return _record_tesseract_path(path_candidate)
        candidate_exe = path_candidate / 'tesseract.exe'
        if candidate_exe.exists():
            return _record_tesseract_path(candidate_exe)

    auto_cmd = shutil.which('tesseract')
    if auto_cmd:
        exe_path = Path(auto_cmd)
        return _record_tesseract_path(exe_path)

    print("⚠️  Tesseract executable not found. Set CLM_TESSERACT_PATH or TESSERACT_CMD to the full tesseract.exe path. OCR will be skipped for scanned PDFs until it's installed.")
    return None


TESSERACT_CMD = _configure_tesseract()
TESSERACT_AVAILABLE = TESSERACT_CMD is not None


def _record_ghostscript_path(executable: Path) -> str:
    os.environ['OCRMYPDF_GS'] = str(executable)
    os.environ['GHOSTSCRIPT_CMD'] = str(executable)
    _ensure_dir_on_path(executable.parent)
    print(f"✅ Ghostscript configured: {executable}")
    return str(executable)


def _probe_ghostscript(base_candidate: Path) -> Optional[Path]:
    exe_names = ('gswin64c.exe', 'gswin32c.exe', 'gs.exe', 'gswin64c', 'gswin32c', 'gs')
    try:
        if base_candidate.is_file() and base_candidate.name.lower().startswith('gs'):
            return base_candidate
    except Exception:
        return None

    search_dirs: List[Path] = []
    try:
        if base_candidate.is_dir():
            search_dirs.append(base_candidate)
            search_dirs.append(base_candidate / 'bin')
            for child in base_candidate.iterdir():
                if child.is_dir() and child.name.lower().startswith('gs'):
                    search_dirs.append(child)
                    search_dirs.append(child / 'bin')
    except Exception:
        pass

    for directory in search_dirs:
        if not directory or not directory.exists():
            continue
        for exe_name in exe_names:
            candidate = directory / exe_name
            if candidate.exists():
                return candidate
    return None


def _configure_ghostscript() -> Optional[str]:
    def normalize(value: str) -> Path:
        return Path(value.strip('"')).expanduser()

    env_cmd = os.getenv('OCRMYPDF_GS') or os.getenv('GHOSTSCRIPT_CMD')
    if env_cmd:
        env_candidate = normalize(env_cmd)
        if env_candidate.exists():
            return _record_ghostscript_path(env_candidate)
        print(f"⚠️  OCRMYPDF_GS/GHOSTSCRIPT_CMD was set but not found at {env_candidate}")

    hint_candidates = [DEFAULT_GHOSTSCRIPT_HINT, r"C:\\Program Files\\gs", r"C:\\Program Files (x86)\\gs"]
    for hint in hint_candidates:
        if not hint:
            continue
        candidate_path = normalize(hint)
        located = _probe_ghostscript(candidate_path)
        if located:
            return _record_ghostscript_path(located)

    for exe_name in ('gswin64c', 'gswin32c', 'gs'):
        auto_cmd = shutil.which(exe_name)
        if auto_cmd:
            return _record_ghostscript_path(Path(auto_cmd))

    print("⚠️  Ghostscript executable (gswin64c.exe) not found. Install it or set CLM_GHOSTSCRIPT_PATH. OCR for scanned PDFs will be skipped until then.")
    return None


GHOSTSCRIPT_CMD = _configure_ghostscript()
GHOSTSCRIPT_AVAILABLE = GHOSTSCRIPT_CMD is not None


# ============================================================================
# CHROMADB CONFIGURATION
# ============================================================================

class ContractVectorDB:
    """Manages ChromaDB for contract storage and semantic search."""
    
    def __init__(self, persist_directory: str = "./chroma_db"):
        if not CHROMADB_AVAILABLE:
            raise RuntimeError("ChromaDB is required. Install with: pip install chromadb")
        
        self.persist_directory = persist_directory
        self.client = chromadb.PersistentClient(path=persist_directory)
        self.collection = self.client.get_or_create_collection(
            name="contracts",
            metadata={"hnsw:space": "cosine"}
        )
        self.embedding_model: Optional[SentenceTransformer] = None
        self.embedding_model_name: Optional[str] = None
        self.embedding_dimension: Optional[int] = None
        self._existing_embedding_dim = self._detect_existing_embedding_dim()

        if EMBEDDINGS_AVAILABLE:
            self._load_embedding_model()
        else:
            print("⚠️  sentence-transformers not available; semantic search disabled")

    def _detect_existing_embedding_dim(self) -> Optional[int]:
        """Return the stored embedding dimension (if any) from the collection."""
        try:
            snapshot = self.collection.get(include=['embeddings'], limit=1)
            embeddings_sample = snapshot.get('embeddings')
            if embeddings_sample is not None and len(embeddings_sample) > 0:
                first_embedding = embeddings_sample[0]
                if first_embedding is not None:
                    dim = len(first_embedding)
                    if dim > 0:
                        print(f"ℹ️  Existing ChromaDB embeddings detected with dimension {dim}.")
                        return dim
        except Exception as err:
            print(f"⚠️  Could not inspect existing embeddings: {err}")
        return None

    def _load_embedding_model(self) -> None:
        """Load an embedding model that is compatible with stored vectors."""

        if EMBEDDING_MODEL_OVERRIDE:
            preferred_models = [EMBEDDING_MODEL_OVERRIDE]
        else:
            preferred_models = [
                "nlpaueb/bert-base-uncased-contracts",
                "nlpaueb/legal-bert-base-uncased",
                "all-MiniLM-L6-v2",
            ]

        # Avoid duplicates while preserving order
        seen = set()
        ordered_models: List[str] = []
        for name in preferred_models:
            if name and name not in seen:
                ordered_models.append(name)
                seen.add(name)

        for model_name in ordered_models:
            try:
                candidate = SentenceTransformer(model_name)
                dimension = candidate.get_sentence_embedding_dimension()
                if self._existing_embedding_dim and dimension != self._existing_embedding_dim:
                    print(
                        f"⚠️  Model '{model_name}' outputs {dimension} dims but existing vectors are "
                        f"{self._existing_embedding_dim} dims. Skipping."
                    )
                    continue
                self.embedding_model = candidate
                self.embedding_model_name = model_name
                self.embedding_dimension = dimension
                if self._existing_embedding_dim:
                    print(
                        f"✅ Embedding model '{model_name}' loaded (dimension {dimension}) to match existing data."
                    )
                else:
                    print(
                        f"✅ Embedding model '{model_name}' loaded (dimension {dimension})."
                    )
                return
            except Exception as err:
                print(f"⚠️  Could not load embedding model '{model_name}': {err}")

        print("❌ No compatible embedding model could be loaded; semantic search disabled.")
        self.embedding_model = None
        self.embedding_model_name = None
        self.embedding_dimension = None

    def contract_exists(self, contract_id: str) -> bool:
        """Return True if the contract ID already exists in the collection."""
        try:
            result = self.collection.get(ids=[contract_id])
            return bool(result.get('ids'))
        except Exception as err:
            print(f"⚠️  Could not check contract '{contract_id}': {err}")
            return False

    def delete_contract(self, contract_id: str) -> None:
        """Remove a contract from the collection (used for force reprocessing)."""
        try:
            self.collection.delete(ids=[contract_id])
        except Exception as err:
            print(f"⚠️  Could not delete contract '{contract_id}': {err}")

    def add_contract(
        self,
        contract_id: str,
        text: str,
        metadata: Dict[str, Any]
    ) -> bool:
        """Add a contract document plus metadata to ChromaDB."""
        if not text:
            print(f"⚠️  Skipping empty document for {contract_id}")
            return False

        # Ensure metadata keys/values are supported types for ChromaDB storage
        safe_metadata = _sanitize_metadata_dict(metadata or {})

        try:
            if self.embedding_model:
                embedding = self.embedding_model.encode(text).tolist()
                self.collection.add(
                    ids=[contract_id],
                    embeddings=[embedding],
                    documents=[text],
                    metadatas=[safe_metadata]
                )
            else:
                self.collection.add(
                    ids=[contract_id],
                    documents=[text],
                    metadatas=[safe_metadata]
                )

            print(f"✅ Contract {contract_id} added to ChromaDB")
            return True
        except Exception as e:
            print(f"❌ Error adding contract: {e}")
            return False
    
    def search_similar_contracts(
        self, 
        query_text: str, 
        n_results: int = 5
    ) -> List[Dict[str, Any]]:
        """
        HYBRID SEARCH: Checks metadata (ID, parties) first, then semantic similarity.
        This ensures exact matches (like 'Keystone' in filename) rank highest.
        """
        if not self.embedding_model:
            print("⚠️  Semantic search unavailable without embeddings")
            return []
        
        try:
            query_lower = query_text.lower()
            query_words = query_lower.split()
            all_contracts = self.collection.get(include=['documents', 'metadatas'])
            scored_contracts = []
            query_embedding = self.embedding_model.encode(query_text).tolist()
            semantic_results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=len(all_contracts['ids'])
            )
            distance_map = {
                id_: dist 
                for id_, dist in zip(semantic_results['ids'][0], semantic_results['distances'][0])
            }
            for idx, contract_id in enumerate(all_contracts['ids']):
                metadata = all_contracts['metadatas'][idx]
                doc_text = all_contracts['documents'][idx]
                metadata_score = 0
                id_lower = contract_id.lower()
                for word in query_words:
                    if len(word) > 2 and word in id_lower:
                        metadata_score += 50
                parties = metadata.get('parties', 'Unknown').lower()
                if parties != 'unknown':
                    for word in query_words:
                        if len(word) > 2 and word in parties:
                            metadata_score += 30
                contract_type = metadata.get('contract_type', '').lower()
                for word in query_words:
                    if len(word) > 2 and word in contract_type:
                        metadata_score += 20
                metadata_score = min(metadata_score, 100)
                distance = distance_map.get(contract_id, 2.0)
                semantic_score = (1 - distance) * 100
                hybrid_score = (metadata_score * 0.7) + (semantic_score * 0.3)
                scored_contracts.append({
                    'id': contract_id,
                    'text': doc_text,
                    'metadata': metadata,
                    'distance': distance,
                    'metadata_score': metadata_score,
                    'semantic_score': semantic_score,
                    'hybrid_score': hybrid_score
                })
            scored_contracts.sort(key=lambda x: x['hybrid_score'], reverse=True)
            return scored_contracts[:n_results]
            
        except Exception as e:
            print(f"❌ Search error: {e}")
            return []
    
    def get_contract(self, contract_id: str) -> Optional[Dict[str, Any]]:
        """Retrieve specific contract by ID."""
        try:
            result = self.collection.get(ids=[contract_id])
            if result['ids']:
                return {
                    'id': result['ids'][0],
                    'text': result['documents'][0],
                    'metadata': result['metadatas'][0]
                }
            return None
        except Exception as e:
            print(f"❌ Error retrieving contract: {e}")
            return None
    
    def get_all_metadata(self) -> List[Dict[str, Any]]:
        """Get metadata for all contracts (for SharePoint list population)."""
        try:
            # ChromaDB doesn't have a direct "get all" - query with high limit
            all_items = self.collection.get()
            metadata_list = []
            for i, id_ in enumerate(all_items['ids']):
                metadata_list.append({
                    'contract_id': id_,
                    **all_items['metadatas'][i]
                })
            return metadata_list
        except Exception as e:
            print(f"❌ Error fetching metadata: {e}")
            return []
    
    def count_contracts(self) -> int:
        """Get total number of contracts in database."""
        return self.collection.count()

# ============================================================================
# DOCUMENT PROCESSING (PDF + DOCX with OCR)
# ============================================================================

def extract_text_from_pdf(
    pdf_path: str,
    use_ocr: bool = False,
    allow_auto_ocr: bool = True
) -> Dict[str, Any]:
    """
    Extract text from PDF using pymupdf4llm (best for LLMs) or PyMuPDF with optional OCR.
    Returns text, tables, and metadata.
    """
    if not PDF_SUPPORT:
        raise RuntimeError("PyMuPDF is required. Install with: pip install pymupdf")
    
    try:
        result = {
            "text": "",
            "tables": [],
            "metadata": {},
            "page_count": 0,
            "has_images": False,
            "ocr_applied": False
        }
        
        # Use pymupdf4llm if available (optimized for LLMs with better layout analysis)
        if PYMUPDF4LLM_AVAILABLE:
            print("📄 Using pymupdf4llm for optimized text extraction...")
            md_text = pymupdf4llm.to_markdown(pdf_path)
            result["text"] = md_text
            
            # Still need fitz for metadata
            doc = fitz.open(pdf_path)
            result["page_count"] = len(doc)
            result["metadata"] = doc.metadata or {}
            
            # Check for images
            for page in doc:
                if page.get_images():
                    result["has_images"] = True
                    break
            doc.close()
        else:
            # Fallback to standard PyMuPDF
            doc = fitz.open(pdf_path)
            result["page_count"] = len(doc)
            result["metadata"] = doc.metadata or {}
            
            for page_num, page in enumerate(doc):
                # Extract text
                page_text = page.get_text()
                result["text"] += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                # Extract tables (PyMuPDF can detect table structures)
                tables = page.find_tables()
                if tables:
                    for table_num, table in enumerate(tables):
                        result["tables"].append({
                            "page": page_num + 1,
                            "table_num": table_num + 1,
                            "data": table.extract()
                        })
                
                # Check for images (potential scanned document)
                images = page.get_images()
                if images:
                    result["has_images"] = True
            
            doc.close()
        
        # If text is sparse but has images, suggest OCR
        stripped_len = len(result["text"].strip())
        should_force_ocr = (
            allow_auto_ocr
            and OCR_SUPPORT
            and TESSERACT_AVAILABLE
            and stripped_len < 400
            and (use_ocr or result["has_images"] or stripped_len == 0)
        )

        if should_force_ocr:
            print(f"📷 Scanned/low-text PDF detected, running OCR...")
            ocr_text = extract_text_with_ocr(pdf_path)
            if ocr_text:
                result["text"] = ocr_text
                result.setdefault("metadata", {})["ocr_applied"] = True
                result["ocr_applied"] = True
        elif use_ocr and not OCR_SUPPORT:
            print("⚠️  OCR requested but ocrmypdf is not installed")
        elif use_ocr and not TESSERACT_AVAILABLE:
            print("⚠️  OCR requested but Tesseract executable is not available")
        elif use_ocr and not GHOSTSCRIPT_AVAILABLE:
            print("⚠️  OCR requested but Ghostscript executable is not available")
        
        return result
    
    except Exception as e:
        print(f"❌ Error extracting PDF: {e}")
        return {"text": "", "tables": [], "metadata": {}, "page_count": 0, "has_images": False}


def extract_text_with_ocr(pdf_path: str) -> str:
    """Extract text from scanned PDF using OCR."""
    if not OCR_SUPPORT:
        print("⚠️  OCR not available. Install with: pip install ocrmypdf")
        return ""
    if not TESSERACT_AVAILABLE:
        print("⚠️  OCR skipped because Tesseract executable is not configured")
        return ""
    if not GHOSTSCRIPT_AVAILABLE:
        print("⚠️  OCR skipped because Ghostscript executable (gswin64c.exe) is not configured")
        return ""
    
    try:
        import tempfile
        
        # Create temporary output file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_output:
            temp_output_path = temp_output.name
        
        # Run OCR
        print(f"🔍 Running OCR on {Path(pdf_path).name}...")
        ocrmypdf.ocr(
            pdf_path,
            temp_output_path,
            force_ocr=True,
            skip_text=False,
            tesseract_timeout=180
        )
        
        # Extract text from OCR'd PDF
        result = extract_text_from_pdf(
            temp_output_path,
            use_ocr=False,
            allow_auto_ocr=False
        )
        
        # Cleanup
        os.unlink(temp_output_path)
        
        print(f"✅ OCR complete: {len(result['text'])} characters extracted")
        return result['text']
    
    except Exception as e:
        print(f"❌ OCR error: {e}")
        return ""


def extract_text_from_docx(docx_path: str) -> Dict[str, Any]:
    """Extract text and tables from Word document."""
    if not DOCX_SUPPORT:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")
    
    try:
        result = {
            "text": "",
            "tables": [],
            "metadata": {}
        }
        
        # Check if file is a OneDrive placeholder or temp file
        file_path = Path(docx_path)
        if file_path.name.startswith('~$') or file_path.name.startswith('～$'):
            print(f"  ⚠️ Skipping Word temp/lock file: {file_path.name}")
            return {"text": "", "tables": [], "metadata": {}, "skipped": True}
        
        # Check file size - very small files may be placeholders
        try:
            file_size = file_path.stat().st_size
            if file_size < 500:
                print(f"  ⚠️ File too small ({file_size} bytes), likely a placeholder")
                return {"text": "", "tables": [], "metadata": {}, "skipped": True}
        except Exception:
            pass
        
        doc = Document(docx_path)
        
        # Extract core properties
        core_props = doc.core_properties
        result["metadata"] = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else ""
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                result["text"] += para.text + "\n"
        
        # Extract tables
        for table_num, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            result["tables"].append({
                "table_num": table_num + 1,
                "data": table_data
            })
        
        return result
    
    except Exception as e:
        print(f"❌ Error extracting DOCX: {e}")
        return {"text": "", "tables": [], "metadata": {}}


# ============================================================================
# UTILITY FUNCTIONS (from original system)
# ============================================================================

def extract_defined_terms(text: str) -> List[Dict[str, str]]:
    """Extract defined terms in format: Term ("Defined Term") or ("the Defined Term")."""
    pattern = r'([^()"]+?)\s*\((?:the\s+)?"([^"]+)"\)'
    matches = re.findall(pattern, text)
    return [{"full_term": match[0].strip(), "short_term": match[1].strip()} for match in matches]


def find_contract_sections(text: str) -> Dict[str, str]:
    """Identify key contract sections."""
    sections = {}
    section_patterns = {
        "preamble": r"(THIS AGREEMENT.{0,100}WHEREAS)",
        "term_and_termination": r"(TERM AND TERMINATION|Term of Agreement)",
        "payment_terms": r"(PAYMENT|Compensation|Fees)",
        "confidentiality": r"(CONFIDENTIALITY|Non-Disclosure)",
        "signatures": r"(IN WITNESS WHEREOF|EXECUTED|Signature:)"
    }
    
    for section_name, pattern in section_patterns.items():
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            sections[section_name] = match.group(1)[:200]  # First 200 chars
    
    return sections


def extract_renewal_terms(text: str) -> Dict[str, Any]:
    """Extract renewal and notice requirements."""
    renewal_info = {
        "has_auto_renewal": bool(re.search(r"automatically renew", text, re.IGNORECASE)),
        "notice_period": None,
        "renewal_term": None
    }
    
    # Look for notice periods (30, 60, 90 days)
    notice_match = re.search(r"(\d+)\s*days?\s*(?:prior\s*)?notice", text, re.IGNORECASE)
    if notice_match:
        renewal_info["notice_period"] = f"{notice_match.group(1)} days"
    
    # Look for renewal term
    term_match = re.search(r"renew for (?:an additional |a )?(\d+[- ](year|month|day))", text, re.IGNORECASE)
    if term_match:
        renewal_info["renewal_term"] = term_match.group(1)
    
    return renewal_info


def detect_pii(text: str) -> List[str]:
    """Detect potential PII in text."""
    pii_patterns = {
        "email": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        "phone": r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
        "ssn": r'\b\d{3}-\d{2}-\d{4}\b',
        "credit_card": r'\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b'
    }
    
    detected = []
    for pii_type, pattern in pii_patterns.items():
        if re.search(pattern, text):
            detected.append(pii_type)
    
    return detected


# Common party role identifiers shared across extraction helpers
PARTY_ROLE_TERMS = [
    'Distributor', 'Supplier', 'Company', 'Customer', 'OEM', 'Licensor', 'Licensee',
    'Vendor', 'Reseller', 'Partner', 'Manufacturer', 'Seller', 'Buyer',
    'Provider', 'Client', 'Contractor', 'Subcontractor'
]
PARTY_ROLE_TERMS_LOWER = {term.lower() for term in PARTY_ROLE_TERMS}
COUNTERPARTY_ROLE_HINTS = {
    'distributor', 'customer', 'licensee', 'reseller', 'partner', 'buyer', 'client',
    'contractor', 'subcontractor', 'oem'
}

# Spectralink addresses for date-based detection
SPECTRALINK_ADDRESSES = {
    'current': {
        'address': '305 S. Arthur Ave. Louisville, CO 80027',
        'effective_from': '2024-12-01',  # December 2024 onwards
        'variants': ['spectralink', 'spectralink corporation', 'spectralink corp']
    },
    'boulder': {
        'address': '2560 55th Street, Boulder, Colorado 80301',
        'effective_until': '2024-11-30',  # Before December 2024
        'variants': ['spectralink', 'spectralink corporation', 'spectralink corp']
    },
    'emea': {
        'address': 'Bygholm Soepark 21 E, 8700 Horsens, Denmark',
        'variants': ['spectralink europe', 'spectralink aps', 'spectralink europe aps']
    }
}


def detect_spectralink_address(company_name: str, effective_date: str, text: str) -> Optional[str]:
    """
    Detect which Spectralink address to use based on company name, date, and text.
    
    Args:
        company_name: Company name from extraction
        effective_date: Effective date in MM/DD/YYYY format or 'Unknown'
        text: Full contract text to search for address hints
    
    Returns:
        Address string if Spectralink detected, None otherwise
    """
    if not company_name:
        return None
    
    company_lower = company_name.lower()
    
    # Check if it's a Spectralink variant - check EMEA first (more specific)
    is_spectralink = False
    location = None
    
    # Check EMEA first (most specific)
    for variant in SPECTRALINK_ADDRESSES['emea']['variants']:
        if variant in company_lower:
            return SPECTRALINK_ADDRESSES['emea']['address']
    
    # Then check US variants
    for loc in ['current', 'boulder']:
        info = SPECTRALINK_ADDRESSES[loc]
        for variant in info['variants']:
            if variant in company_lower:
                is_spectralink = True
                location = loc
                break
        if is_spectralink:
            break
    
    if not is_spectralink:
        return None
    
    # For US Spectralink, determine based on date
    if effective_date and effective_date != 'Unknown':
        try:
            # Parse MM/DD/YYYY
            from datetime import datetime
            date_obj = datetime.strptime(effective_date, '%m/%d/%Y')
            cutoff_date = datetime.strptime('12/01/2024', '%m/%d/%Y')
            
            if date_obj >= cutoff_date:
                return SPECTRALINK_ADDRESSES['current']['address']
            else:
                return SPECTRALINK_ADDRESSES['boulder']['address']
        except ValueError:
            pass
    
    # Fallback: search text for address hints
    if '305 S. Arthur' in text or 'Louisville, CO' in text:
        return SPECTRALINK_ADDRESSES['current']['address']
    elif '2560 55th' in text or 'Boulder, Colorado 80301' in text:
        return SPECTRALINK_ADDRESSES['boulder']['address']
    elif 'Bygholm' in text or 'Horsens, Denmark' in text:
        return SPECTRALINK_ADDRESSES['emea']['address']
    
    # Default to current address if date unknown
    return SPECTRALINK_ADDRESSES['current']['address']


EU_ENTITY_PATTERNS = [
    r"\b[A-Z][A-Za-z&\s]+ Limited\b",
    r"\b[A-Z][A-Za-z&\s]+ Ltd\.?\b",
    r"\b[A-Z][A-Za-z&\s]+ GmbH\b",
    r"\b[A-Z][A-Za-z&\s]+ ApS\b",
    r"\b[A-Z][A-Za-z&\s]+ BV\b",
    r"\b[A-Z][A-Za-z&\s]+ NV\b",
    r"\b[A-Z][A-Za-z&\s]+ AG\b",
    r"\b[A-Z][A-Za-z&\s]+ SA\b",
    r"\b[A-Z][A-Za-z&\s]+ SAS\b"
]

LEGAL_ENTITY_KEYWORDS = {
    'corporation', 'corp', 'llc', 'inc', 'ltd', 'limited', 'gmbh', 'aps', 'sa', 'sas',
    'bv', 'nv', 'ag', 'company', 'limited liability', 'pte', 'pty', 'co.', 'sarl', 'oy',
    'oyj', 'ab', 'kk', 'pte. ltd', 'pty ltd'
}
PARTY_SNIPPET_STOPWORDS = {
    'agreement', 'clause', 'section', 'article', 'warranty', 'provision', 'subject',
    'hereof', 'hereto', 'thereof', 'deliverable', 'services', 'term', 'notice', 'parties'
}

FUZZY_DUPLICATE_THRESHOLD = 92

PATH_STATUS_KEYWORDS = {
    'Terminated': ['terminat'],
    'Draft': ['draft'],
    'Archive': ['archive']
}

ROLE_KEYWORDS = [
    ('Statement of Work', ['statement of work', 'sow']),
    ('Amendment/Addendum', ['amendment', 'addendum', 'appendix', 'schedule', 'exhibit']),
    ('Termination Letter', ['termination letter', 'notice of termination']),
    ('Summary/Checklist', ['summary', 'checklist', 'processing checklist']),
    ('Contract Summary', ['contract summary']),
    ('Master Agreement', ['master services agreement', 'master agreement', 'msa'])
]

_FOLDER_CHILD_COUNT_CACHE: Dict[str, int] = {}


def _dedupe_candidates(candidates: List[str]) -> List[str]:
    """Remove duplicate party names using fuzzy matching when available."""
    unique: List[str] = []
    for candidate in candidates:
        if not candidate:
            continue
        duplicate = False
        for existing in unique:
            if RAPIDFUZZ_AVAILABLE:
                score = fuzz.token_set_ratio(candidate, existing)
                if score >= FUZZY_DUPLICATE_THRESHOLD:
                    duplicate = True
                    break
            else:
                if candidate.lower() == existing.lower():
                    duplicate = True
                    break
        if not duplicate:
            unique.append(candidate)
    return unique


def _prioritize_primary_company(candidates: List[str]) -> List[str]:
    """Move configured primary company names to the front of the list."""
    if not PRIMARY_COMPANY_NAMES or not candidates:
        return candidates

    best_idx = None
    best_score = 0
    primary_list = list(PRIMARY_COMPANY_NAMES)

    for idx, candidate in enumerate(candidates):
        candidate_lower = candidate.lower()
        match_score = 0
        if RAPIDFUZZ_AVAILABLE:
            match = process.extractOne(candidate, primary_list, scorer=fuzz.token_set_ratio)
            if match:
                match_score = match[1]
        else:
            match_score = 100 if any(name in candidate_lower for name in primary_list) else 0

        if match_score > best_score:
            best_idx = idx
            best_score = match_score

    if best_idx is not None and best_score >= 80:
        candidate = candidates.pop(best_idx)
        candidates.insert(0, candidate)

    return candidates


def _ensure_standard_metadata(metadata: Dict[str, Any]) -> None:
    """Backfill any missing standard metadata fields with defaults."""
    for field, default in STANDARD_METADATA_DEFAULTS.items():
        value = metadata.get(field)
        if value in (None, '', []):
            metadata[field] = default if default not in (None, '') else 'Unknown'

EMEA_COUNTRIES = [
    'Denmark', 'Germany', 'France', 'UK', 'United Kingdom', 'England', 'Scotland', 'Netherlands',
    'Belgium', 'Sweden', 'Norway', 'Finland', 'Ireland', 'Spain', 'Italy', 'Austria', 'Switzerland',
    'Poland', 'Czech Republic', 'Portugal', 'Greece', 'Hungary', 'Romania', 'Bulgaria', 'Croatia',
    'Slovakia', 'Lithuania', 'Latvia', 'Estonia', 'Iceland', 'Luxembourg', 'Liechtenstein', 'Malta'
]
NA_LOCATIONS = [
    'United States', 'USA', 'U.S.', 'Canada', 'Mexico', 'Colorado', 'California', 'Texas',
    'New York', 'Delaware', 'Florida', 'Illinois', 'Massachusetts', 'Washington', 'Virginia'
]
APAC_COUNTRIES = [
    'China', 'Japan', 'Korea', 'Singapore', 'Australia', 'India', 'Thailand', 'Malaysia', 'Indonesia',
    'Philippines', 'Vietnam', 'Taiwan', 'Hong Kong', 'New Zealand'
]

DISTRIBUTION_KEYWORDS = ['distributor', 'distribution', 'reseller', 'partner program']
PROCUREMENT_KEYWORDS = ['supplier', 'manufactur', 'oem', 'vendor', 'hotel', 'credit card', 'procurement']
MARKETING_KEYWORDS = ['marketing', 'campaign', 'advertising', 'joint marketing']


def _is_ocr_garbage(text: str) -> bool:
    """Check if text contains OCR garbage artifacts that make it unusable."""
    if not text or len(text) < 2:
        return True
    
    # Check for common OCR garbage patterns
    garbage_patterns = [
        r'#{2,}',           # Multiple hash marks (common OCR artifact)
        r'`{2,}',           # Multiple backticks
        r'\\[A-Za-z]',      # File path-like content  
        r'\d{1,2}\s+[A-Z][a-z]+\s+\d{4}\s+##',  # Date followed by hashes
        r'[^\x00-\x7F]{3,}', # Multiple non-ASCII chars in a row (encoding issues)
        r'\*{3,}',          # Multiple asterisks
        r'_{3,}',           # Multiple underscores
        r'\.{5,}',          # Many dots in a row
    ]
    
    for pattern in garbage_patterns:
        if re.search(pattern, text):
            return True
    
    # Check if text has too many digits/symbols vs letters (likely garbage)
    letters = sum(1 for c in text if c.isalpha())
    digits_symbols = sum(1 for c in text if c.isdigit() or c in '#@$%^&*_+={}[]|\\')
    if len(text) > 10 and digits_symbols > letters:
        return True
    
    return False


def _normalize_clause_text(clause: str) -> str:
    clause = re.sub(r'\s+', ' ', clause).strip(' ,;')
    clause = re.sub(r'^and\s+', '', clause, flags=re.IGNORECASE)
    clause = re.sub(r'\((?:the\s+)?["\u201C\u201D]?[A-Za-z\s&-]+["\u201C\u201D]?\)\s*$', '', clause)
    
    # OCR noise cleanup - remove common OCR artifacts
    clause = re.sub(r'#{2,}', '', clause)  # Multiple hash marks
    clause = re.sub(r'`+', '', clause)  # Backticks
    clause = re.sub(r'<[A-Z]\\[^>]*>', '', clause)  # File path artifacts
    clause = re.sub(r'\\[A-Za-z][^\\]*\\', ' ', clause)  # Path separators
    clause = re.sub(r'\s+', ' ', clause).strip()  # Re-normalize whitespace
    
    return clause


def _is_primary_company_clause(clause: str) -> bool:
    """Return True when the clause references the configured primary company."""
    lower_clause = clause.lower()
    if PRIMARY_COMPANY_NAMES and any(name in lower_clause for name in PRIMARY_COMPANY_NAMES):
        return True
    if PRIMARY_COMPANY_ADDRESSES and any(addr in lower_clause for addr in PRIMARY_COMPANY_ADDRESSES):
        return True
    return False


def _split_name_address(clause: str) -> Dict[str, str]:
    """Return separate name/address strings when possible."""
    name = clause
    address = ""
    pattern = (
        r'^(?P<name>.+?)(?:,?\s+(?:a|an)\s+[\w\s]+?)?(?:,?\s+(?:having\s+(?:its\s+)?principal\s+place\s+of\s+business\s+at|'
        r'located\s+at|with\s+offices\s+at|with\s+its\s+registered\s+office\s+at|at))\s+(?P<address>.+)$'
    )
    match = re.search(pattern, clause, re.IGNORECASE)
    if match:
        name = match.group('name').strip(' ,;')
        address = match.group('address').strip(' ,;')
    else:
        clause_lower = clause.lower()
        for addr in PRIMARY_COMPANY_ADDRESSES:
            idx = clause_lower.find(addr)
            if idx != -1:
                name = clause[:idx].strip(' ,;')
                address = clause[idx: idx + len(addr)]
                break
    if address:
        address = re.sub(r'\s+', ' ', address)
        address = re.split(r'\s*\(["\u201C\u201D]|"', address)[0].strip(' ,;')
    
    # Clean up name - remove OCR artifacts and truncate if too long
    name = re.sub(r'#{2,}|`+', '', name).strip()
    if len(name) > 150:
        # Try to find the actual company name (first legal entity)
        legal_match = re.search(r'^([A-Z][A-Za-z0-9&.,\s-]+?(?:Corporation|Corp\.?|Company|LLC|Limited|Ltd\.?|GmbH|ApS|AG|BV|Inc\.?|AB|A/S))', name, re.IGNORECASE)
        if legal_match:
            name = legal_match.group(1).strip()
        else:
            name = name[:150].rsplit(' ', 1)[0]  # Truncate at word boundary
    
    return {'name': name, 'address': address}


def get_party_clauses(text: str) -> Dict[str, Optional[str]]:
    """Extract party clauses from the preamble for reuse."""
    result: Dict[str, Optional[str]] = {
        'primary_party': None,
        'counterparty': None,
        'ordered': [],
        'primary_party_info': None,
        'counterparty_info': None,
        'ordered_info': []
    }
    preamble = text
    pattern = r'(?:between|by\s+and\s+between)\s+(.{30,4000}?)\s+(?:and|&)\s+(.{30,4000}?)(?:\.|\n{2}|;|\r)'
    match = re.search(pattern, preamble, re.IGNORECASE | re.DOTALL)
    if not match:
        alt_pattern = r'This\s+[A-Za-z\s]*?Agreement[\s\S]{0,2000}?(?:between|by\s+and\s+between)\s+(.{30,4000}?)\s+(?:and|&)\s+(.{30,4000}?)(?:\.|\n{2}|;|\r)'
        match = re.search(alt_pattern, preamble, re.IGNORECASE | re.DOTALL)
    if not match:
        return result
    clauses = [_normalize_clause_text(match.group(1)), _normalize_clause_text(match.group(2))]
    ordered = [clause for clause in clauses if clause]
    result['ordered'] = ordered
    for clause in clauses:
        if not clause:
            continue
        info = _split_name_address(clause)
        result['ordered_info'].append(info)
        if _is_primary_company_clause(clause):
            result['primary_party'] = clause
            result['primary_party_info'] = info
        elif not result['counterparty']:
            result['counterparty'] = clause
            result['counterparty_info'] = info
    return result


def _clean_party_candidate(candidate: Optional[str]) -> Optional[str]:
    if not candidate:
        return None
    cleaned = re.sub(r'\s+', ' ', candidate).strip(' "\'.,;:-')
    if not cleaned:
        return None
    
    # OCR noise cleanup - reject candidates with obvious garbage
    # Check for OCR artifacts: multiple #, backticks, file paths, random punctuation sequences
    if re.search(r'#{2,}|`{1,}|\\[A-Za-z]|<[A-Z]\\|[#]{3,}', cleaned):
        return None
    
    # Reject if it looks like a file path
    if re.search(r'\\[A-Za-z0-9_-]+\\', cleaned):
        return None
    
    # Reject if more than 30% non-alphanumeric characters (sign of OCR noise)
    alphanumeric = sum(1 for c in cleaned if c.isalnum() or c.isspace())
    if len(cleaned) > 10 and alphanumeric / len(cleaned) < 0.7:
        return None
    
    # Reject if starts with numbers only (likely address fragment)
    if re.match(r'^\d{4,}\s+[A-Z]', cleaned):
        return None
    
    # honor request to avoid aggressive truncation; only trim extremely long blobs
    if len(cleaned) > 200:
        # For party names, 200 chars is plenty - truncate aggressively for OCR docs
        trimmed = cleaned[:200]
        last_space = trimmed.rfind(' ')
        cleaned = trimmed if last_space == -1 else trimmed[:last_space]
    return cleaned


def _looks_like_party_name(candidate: str) -> bool:
    lower_candidate = candidate.lower()
    if any(keyword in lower_candidate for keyword in LEGAL_ENTITY_KEYWORDS):
        return True
    words = candidate.split()
    if 1 <= len(words) <= 4 and all(word[:1].isupper() for word in words if word):
        return True
    return False


def _normalize_party_candidate(candidate: Optional[str]) -> Optional[str]:
    cleaned = _clean_party_candidate(candidate)
    if not cleaned:
        return None
    lower_candidate = cleaned.lower()
    if any(stopword in lower_candidate for stopword in PARTY_SNIPPET_STOPWORDS) and not any(keyword in lower_candidate for keyword in LEGAL_ENTITY_KEYWORDS):
        return None
    if _looks_like_party_name(cleaned):
        return cleaned
    return None


def _address_mentions(address: str, keywords: List[str]) -> bool:
    lower_address = address.lower()
    return any(keyword.lower() in lower_address for keyword in keywords)


def _infer_region_from_addresses(party_details: Dict[str, str]) -> Optional[str]:
    addresses = [value for key, value in party_details.items() if key.endswith('address') and value]
    for address in addresses:
        if _address_mentions(address, EMEA_COUNTRIES):
            return 'EMEA'
    for address in addresses:
        if _address_mentions(address, NA_LOCATIONS):
            return 'North America'
    for address in addresses:
        if _address_mentions(address, APAC_COUNTRIES):
            return 'APAC'
    return None


def assign_contract_owners(contract_type: str, region: Optional[str]) -> Dict[str, Any]:
    owners = _filter_owners(OWNER_LEGAL)
    department = 'Legal'
    lower_type = (contract_type or '').lower()
    normalized_region = (region or '').lower()

    def set_distribution(region_value: Optional[str]):
        normalized = (region_value or '').lower()
        if normalized in {'emea', 'europe', 'europe, middle east and africa'}:
            local_department = 'EMEA Distribution'
            local_owners = _filter_owners(OWNER_DISTRIBUTION_EMEA, OWNER_LEGAL)
        elif normalized in {'north america', 'na', 'apac', 'asia pac', 'asia-pacific'}:
            local_department = 'NA APAC Distributor Management'
            local_owners = _filter_owners(OWNER_DISTRIBUTION_NA_APAC, OWNER_LEGAL)
        else:
            local_department = 'NA Distribution'
            local_owners = _filter_owners(OWNER_DISTRIBUTION_NA, OWNER_LEGAL)
        return local_owners or _filter_owners(OWNER_LEGAL), local_department

    if any(keyword in lower_type for keyword in MARKETING_KEYWORDS):
        owners = _filter_owners(OWNER_LEGAL, OWNER_MARKETING)
        department = 'Marketing'
    elif any(keyword in lower_type for keyword in PROCUREMENT_KEYWORDS):
        owners = _filter_owners(OWNER_LEGAL, OWNER_PROCUREMENT)
        department = 'Procurement'
    elif any(keyword in lower_type for keyword in DISTRIBUTION_KEYWORDS):
        owners, department = set_distribution(normalized_region)

    owners = _ensure_owner_defaults(owners)

    return {
        'owners': owners,
        'department': department
    }


def derive_sharepoint_metadata(contract_type: str, party_details: Dict[str, str], gdpr_info: Dict[str, Any]) -> Dict[str, Any]:
    fields: Dict[str, Any] = {}
    region = gdpr_info.get('region')
    if not region:
        inferred_region = _infer_region_from_addresses(party_details)
        if inferred_region:
            region = inferred_region
    if region:
        fields['region'] = region

    gdpr_flag = 'Yes' if gdpr_info.get('subject_to_gdpr') else 'No'
    gdpr_reason = gdpr_info.get('gdpr_reason', '')
    if gdpr_flag == 'No':
        addresses = [value for key, value in party_details.items() if key.endswith('address') and value]
        for address in addresses:
            if _address_mentions(address, EMEA_COUNTRIES):
                gdpr_flag = 'Yes'
                gdpr_reason = 'Party address in EMEA'
                if not region:
                    region = 'EMEA'
                    fields['region'] = 'EMEA'
                break
    fields['gdpr_applicable'] = gdpr_flag
    if gdpr_reason:
        fields['gdpr_reason'] = gdpr_reason

    owner_info = assign_contract_owners(contract_type, region)
    owner_list = owner_info['owners']
    fields['sharepoint_department'] = owner_info['department']
    if owner_list:
        fields['sharepoint_owner_primary'] = owner_list[0]
        if len(owner_list) > 1:
            fields['sharepoint_owner_secondary'] = owner_list[1]
        fields['sharepoint_owners'] = ', '.join(owner_list)

    return fields


def _folder_child_count(directory: Path) -> int:
    try:
        cache_key = str(directory.resolve())
    except Exception:
        cache_key = str(directory)
    if cache_key in _FOLDER_CHILD_COUNT_CACHE:
        return _FOLDER_CHILD_COUNT_CACHE[cache_key]
    try:
        count = len(list(directory.iterdir()))
    except Exception:
        count = 0
    _FOLDER_CHILD_COUNT_CACHE[cache_key] = count
    return count


def derive_document_context(file_path: str, text: str, relative_label: Optional[str] = None) -> Dict[str, Any]:
    path_obj = Path(file_path)
    folder = path_obj.parent
    parent_folder = folder.parent if folder else None
    grandparent_folder = parent_folder.parent if parent_folder else None
    parts_lower = [part.lower() for part in path_obj.parts]
    statuses = set()

    for label, keywords in PATH_STATUS_KEYWORDS.items():
        if any(any(keyword in part for keyword in keywords) for part in parts_lower):
            statuses.add(label)

    text_window = text.lower() if len(text) <= 60000 else text[:60000].lower()
    if 'notice of termination' in text_window or 'terminated effective' in text_window:
        statuses.add('Terminated')
    if 'draft' in text_window[:5000]:
        statuses.add('Draft')

    if not statuses:
        statuses.add('Active')

    contract_status = 'Terminated' if 'Terminated' in statuses else (
        'Draft' if 'Draft' in statuses else (
            'Archive' if 'Archive' in statuses else 'Active'
        )
    )

    filename_lower = path_obj.name.lower()
    document_role = None
    for label, keywords in ROLE_KEYWORDS:
        if any(keyword in filename_lower for keyword in keywords) or any(keyword in text_window for keyword in keywords):
            document_role = label
            break

    if document_role == 'Statement of Work':
        filename_has_msa = ('msa' in filename_lower) or ('master service agreement' in filename_lower)
        filename_has_sow = 'sow' in filename_lower
        if filename_has_msa and not filename_has_sow:
            document_role = 'Master Agreement'

    relationship_hints = []
    if document_role in {'Statement of Work', 'Amendment/Addendum'}:
        relationship_hints.append('References underlying master agreement')
    if 'master services agreement' in text_window or 'master agreement' in text_window:
        relationship_hints.append('Mentions master services agreement in body')
    if 'schedule' in filename_lower and 'statement of work' in text_window:
        relationship_hints.append('Schedule references SOW terms')

    sequence_label = None
    sequence_token = None
    sow_match = re.search(r'sow[\s_\-]*(\d+)', filename_lower)
    if sow_match:
        sequence_token = sow_match.group(1)
        sequence_label = f"SOW #{sequence_token}"
    elif document_role == 'Statement of Work':
        sow_text = re.search(r'statement of work\s*(?:number|no\.?|#)\s*(\d+)', text_window)
        if sow_text:
            sequence_token = sow_text.group(1)
            sequence_label = f"SOW #{sequence_token}"

    if not sequence_label:
        amendment_match = re.search(r'(?:amendment|addendum)\s*(?:number|no\.?|#)?\s*(\d+)', filename_lower)
        if amendment_match and document_role == 'Amendment/Addendum':
            sequence_token = amendment_match.group(1)
            sequence_label = f"Amendment #{sequence_token}"

    numeric_token = re.search(r'(\d{6,8})', path_obj.stem)
    sequence_date = _format_numeric_token_as_date(numeric_token.group(1)) if numeric_token else None
    if sequence_date:
        relationship_hints.append(f"Sequence date token: {sequence_date}")

    expected_related_documents = []
    if document_role == 'Statement of Work':
        expected_related_documents.append('Master Services Agreement (MSA)')
    if document_role == 'Amendment/Addendum':
        expected_related_documents.append('Underlying Agreement referenced in amendment')
    if document_role == 'Termination Letter':
        expected_related_documents.append('Agreement referenced in termination notice')

    group_source = str(folder).lower()
    contract_group_id = hashlib.md5(group_source.encode('utf-8')).hexdigest()[:12] if group_source else ''
    folder_document_count = _folder_child_count(folder) if folder else 0

    lineage_parts = [part for part in [grandparent_folder.name if grandparent_folder else None,
                                       parent_folder.name if parent_folder else None,
                                       folder.name if folder else None] if part]
    parent_reference = ' / '.join(lineage_parts[-2:]) if lineage_parts else ''

    if parent_reference and document_role in {'Statement of Work', 'Amendment/Addendum', 'Termination Letter'}:
        relationship_hints.append(f"Likely tied to {parent_reference}")

    relationship_text = '; '.join(dict.fromkeys(relationship_hints)) if relationship_hints else None

    return {
        'relative_path': relative_label or str(path_obj),
        'source_folder': str(folder) if folder else '',
        'folder_name': folder.name if folder else '',
        'parent_folder_name': parent_folder.name if parent_folder else '',
        'grandparent_folder_name': grandparent_folder.name if grandparent_folder else '',
        'folder_depth': len(path_obj.parts),
        'folder_document_count': folder_document_count,
        'has_related_documents': folder_document_count > 1,
        'document_status': sorted(statuses),
        'document_status_label': ', '.join(sorted(statuses)),
        'contract_status': contract_status,
        'is_draft': contract_status == 'Draft',
        'is_terminated': contract_status == 'Terminated',
        'is_archive': contract_status == 'Archive',
        'document_role': document_role,
        'expected_related_documents': expected_related_documents,
        'related_contract_hint': relationship_text,
        'contract_group_id': contract_group_id,
        'document_sequence_label': sequence_label,
        'document_sequence_token': sequence_token,
        'document_parent_reference': parent_reference
    }


def detect_contract_type_from_folder(folder_name: str) -> Optional[str]:
    """
    Detect contract type from folder name - more reliable than OCR text.
    Folder names are human-created and accurate.
    
    Common abbreviations handled:
    - Disty/Dist = Distributor
    - Agr/Agmt/Agrmt = Agreement
    - MSA = Master Service Agreement
    - SOW = Statement of Work
    - NDA = Non-Disclosure Agreement
    """
    if not folder_name:
        return None
    
    folder_lower = folder_name.lower()
    
    # Folder-based patterns (order matters - most specific first)
    # Include common abbreviations: agr, agmt, agrmt, dist, disty
    folder_patterns = [
        # Distributor patterns - MOST COMMON, check first with all abbreviations
        (r"distribut(?:or|ion)?\s*(?:agr|agmt|agrmt|agreement)", "Distributor Agreement"),
        (r"dist(?:y|i)?\s*(?:agr|agmt|agrmt|agreement)", "Distributor Agreement"),  # Disty Agr, Dist Agr
        (r"distributor", "Distributor Agreement"),  # Folder just says "Distributor"
        
        # Development agreements
        (r"development\s*(?:agr|agmt|agreement)", "Development Agreement"),
        
        # Reseller/VAR
        (r"reseller\s*(?:agr|agmt|agreement)|var\s*(?:agr|agmt|agreement)", "Reseller Agreement"),
        
        # Partner program - be careful not to match "Service Partner"
        (r"(?<!service\s)partner\s*(?:program\s*)?(?:agr|agmt|agreement)|partnership\s*(?:agr|agmt|agreement)", "Partner Program Agreement"),
        
        # OEM
        (r"oem\s*(?:agr|agmt|agreement)", "OEM Agreement"),
        
        # Sales/Purchase
        (r"(?:sales|purchase|direct\s*sales)\s*(?:agr|agmt|agreement)", "Sales Agreement"),
        
        # SOW - only match explicit SOW, not "service" words
        (r"\bsow\b|statement\s*of\s*work", "Statement of Work (SOW)"),
        
        # MSA
        (r"\bmsa\b|master\s*service(?:\s*agr|\s*agmt|\s*agreement)?", "Master Service Agreement (MSA)"),
        
        # Service Agreement (general) - after MSA
        (r"service\s*(?:agr|agmt|agreement)", "Service Agreement"),
        
        # License
        (r"licens(?:e|ing)\s*(?:agr|agmt|agreement)|software\s*license", "Software License Agreement"),
        
        # SLA
        (r"\bsla\b|service\s*level", "Service Level Agreement (SLA)"),
        
        # Employment
        (r"employment\s*(?:agr|agmt|agreement)|offer\s*letter", "Employment Agreement"),
        
        # Independent Contractor
        (r"independent\s*contractor|contractor\s*(?:agr|agmt|agreement)|freelance", "Independent Contractor Agreement"),
        
        # Consulting
        (r"consulting\s*(?:agr|agmt|agreement)|professional\s*services", "Consulting Agreement"),
        
        # Vendor/Supplier
        (r"vendor\s*(?:agr|agmt|agreement)|supplier\s*(?:agr|agmt|agreement)", "Vendor/Supplier Agreement"),
        
        # Amendment/Addendum - check filename context too
        (r"(?:^|\s)(?:amendment|addendum|amend\s*#?\d*)", "Amendment/Addendum"),
        
        # NDA
        (r"\bnda\b|non[- ]?disclosure|confidentiality", "Non-Disclosure Agreement (NDA)"),
        
        # Terms & Conditions
        (r"terms\s*(?:and|&)\s*conditions|t&c", "Terms & Conditions"),
        
        # Letter of Intent
        (r"letter\s*of\s*intent|\bloi\b", "Letter of Intent"),
        
        # Warranty
        (r"warranty", "Warranty Agreement"),
        
        # Support/Maintenance
        (r"support\s*(?:agr|agmt|agreement)|maintenance", "Support/Maintenance Agreement"),
    ]
    
    for pattern, contract_type in folder_patterns:
        if re.search(pattern, folder_lower):
            return contract_type
    
    return None


def detect_contract_type_from_filename(filename: str) -> Optional[str]:
    """
    Detect contract type from filename when folder doesn't give clear indication.
    Secondary source after folder name.
    """
    if not filename:
        return None
    
    filename_lower = filename.lower()
    
    # Filename patterns - similar to folder but with additional document-specific patterns
    filename_patterns = [
        # Specific document types first
        (r"distribut(?:or|ion)?\s*(?:agr|agmt|agreement)", "Distributor Agreement"),
        (r"dist(?:y|i)?\s*(?:agr|agmt|agreement)", "Distributor Agreement"),
        (r"distributor\s*(?:service\s*)?program", "Distributor Agreement"),  # "Distributor Service Program"
        
        # Reseller Agreement
        (r"reseller\s*(?:agr|agmt|agreement|program)", "Reseller Agreement"),
        (r"service\s*reseller\s*program", "Reseller Agreement"),  # "Service Reseller Program"
        
        # ODM Agreement (Original Design Manufacturing)
        (r"\bodm\b|original\s*design\s*manufacturing", "ODM Agreement"),
        
        # OEM Agreement
        (r"\boem\b|original\s*equipment\s*manufacturer", "OEM Agreement"),
        
        # Amendment patterns
        (r"amend(?:ment)?\s*(?:#|no\.?)?\s*\d+|addendum", "Amendment/Addendum"),
        (r"schedule\s*[a-z]\s*(?:update|amendment)", "Amendment/Addendum"),  # Schedule B Update
        
        # Settlement/Termination
        (r"settlement|termination\s*(?:agr|agreement|letter)", "Termination Agreement"),
        
        # Cover letters, notices - these are correspondence, not contracts
        (r"cover\s*letter|status\s*letter|notice\s*of|consent\s*to", "Correspondence/Notice"),
        
        # Partner program - but not "Service Partner" which is distributor addendum
        (r"(?<!distributor\s)(?<!service\s)partner\s*(?:program|agr)", "Partner Program Agreement"),
        
        # Service Partner is usually a distributor addendum
        (r"service\s*partner\s*(?:support\s*)?program", "Amendment/Addendum"),
        
        # SOW
        (r"\bsow\s*#?\d*\b|statement\s*of\s*work", "Statement of Work (SOW)"),
        
        # MSA
        (r"\bmsa\b|master\s*(?:service\s*)?agreement", "Master Service Agreement (MSA)"),
        
        # Services Agreement (general)
        (r"services?\s*(?:agr|agmt|agreement)", "Service Agreement"),
        
        # Sales/Direct Sales Agreement
        (r"(?:direct\s+)?sales\s*(?:agr|agmt|agreement)", "Sales Agreement"),
        
        # NDA/CDA - Enhanced to catch more variations
        (r"\bnda\b|\bcda\b|non[- ]?disclosure|confidentiality|mutual\s*(?:cda|nda)", "Non-Disclosure Agreement (NDA)"),
    ]
    
    for pattern, contract_type in filename_patterns:
        if re.search(pattern, filename_lower):
            return contract_type
    
    return None


def detect_contract_type(text: str, folder_name: Optional[str] = None, filename: Optional[str] = None) -> str:
    """
    PHASE 1: Quick regex-based contract type detection (fast, deterministic).
    
    Priority order:
    1. Folder name (most reliable - human-created)
    2. Filename (secondary - also human-created)
    3. Document text (fallback - may have OCR issues)
    """
    # PRIORITY 1: Check folder name first (most reliable)
    if folder_name:
        folder_type = detect_contract_type_from_folder(folder_name)
        if folder_type:
            return folder_type
    
    # PRIORITY 2: Check filename
    if filename:
        filename_type = detect_contract_type_from_filename(filename)
        if filename_type:
            return filename_type
    
    # PRIORITY 3: Fall back to text analysis
    text_sample = text.lower()[:10000]  # Only check first 10K chars for speed
    
    # Contract type patterns (order matters - most specific first)
    type_patterns = [
        # Distributor - check multiple variations
        (r"distributor\s+(?:agreement|service\s+program)", "Distributor Agreement"),
        (r"distribution\s+agreement", "Distributor Agreement"),
        (r"distributorship\s+agreement", "Distributor Agreement"),
        
        # Reseller
        (r"reseller\s+agreement|var\s+agreement|value\s+added\s+reseller", "Reseller Agreement"),
        
        # Partner - but not service partner (which is distributor addendum)
        (r"(?<!service\s)partner\s+(?:program\s+)?agreement|partnership\s+agreement", "Partner Program Agreement"),
        
        # OEM
        (r"oem\s+agreement|original\s+equipment\s+manufacturer", "OEM Agreement"),
        
        # Sales/Purchase
        (r"(?:sales|purchase|direct\s+sales)\s+agreement", "Sales Agreement"),
        
        # SOW - be specific
        (r"statement\s+of\s+work|this\s+sow\s+|sow\s+#\d+", "Statement of Work (SOW)"),
        
        # MSA
        (r"master\s+service\s+agreement|this\s+msa\s+", "Master Service Agreement (MSA)"),
        
        # General Service Agreement
        (r"(?<!master\s)service\s+agreement", "Service Agreement"),
        
        # License
        (r"licens(?:e|ing)\s+agreement|software\s+license", "Software License Agreement"),
        
        # SLA
        (r"service\s+level\s+agreement", "Service Level Agreement (SLA)"),
        
        # Employment
        (r"employment\s+agreement|offer\s+letter", "Employment Agreement"),
        
        # Contractor
        (r"independent\s+contractor\s+agreement|contractor\s+agreement", "Independent Contractor Agreement"),
        
        # Consulting
        (r"consulting\s+agreement|professional\s+services\s+agreement", "Consulting Agreement"),
        
        # Vendor
        (r"vendor\s+agreement|supplier\s+agreement", "Vendor/Supplier Agreement"),
        
        # Amendment - check for explicit amendment language
        (r"amendment\s+(?:no\.?\s*)?\d+\s+to|first\s+amendment|second\s+amendment|this\s+amendment", "Amendment/Addendum"),
        (r"addendum\s+to|this\s+addendum", "Amendment/Addendum"),
        
        # NDA
        (r"non[- ]?disclosure\s+agreement|confidentiality\s+agreement|mutual\s+(?:cda|nda)", "Non-Disclosure Agreement (NDA)"),
    ]
    
    for pattern, contract_type in type_patterns:
        if re.search(pattern, text_sample):
            return contract_type
    
    return "Auto-detected"


def _extract_date_near_label(label: str, text: str) -> Optional[str]:
    pattern = re.compile(rf'{re.escape(label)}[\s:,-]{{0,24}}', re.IGNORECASE)
    for match in pattern.finditer(text):
        window_start = max(0, match.start() - 40)
        window_end = min(len(text), match.end() + 200)
        window = text[window_start:window_end]
        literal = _extract_date_literal(window)
        if literal:
            return literal
    return None


def extract_dates(text: str) -> dict:
    """
    PHASE 1: Extract key dates using regex patterns.
    Returns dict with clean date values (not the full phrase).
    """
    dates = {}
    
    # Get defined terms which now have clean values
    defined_terms = extract_defined_terms(text)
    def assign_date(term_label: str, metadata_key: str) -> None:
        raw_value = defined_terms.get(term_label)
        literal_value = _extract_date_literal(raw_value) if raw_value else None
        if not literal_value and raw_value:
            literal_value = _extract_date_near_label(term_label, text)
        if literal_value:
            dates[metadata_key] = literal_value

    assign_date('Effective Date', 'effective_date')
    assign_date('Expiration Date', 'expiration_date')
    assign_date('Termination Date', 'termination_date')
    
    # Fallback patterns for effective date if not found in defined terms
    if 'effective_date' not in dates:
        # Pattern 1: "made effective [date]" or "effective as of [date]"
        effective_patterns = [
            r'(?:made\s+effective|effective\s+as\s+of|effective\s+on|effective\s+date[:\s]+is)\s+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
            r'dated\s+as\s+of\s+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
            r'as\s+of\s+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
        ]
        for pattern in effective_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                dates['effective_date'] = match.group(1)
                break
    
    return dates


def normalize_date_format(date_str: str) -> str:
    """
    Normalize various date formats to MM/DD/YYYY for consistency.
    Handles:
    - October 21, 2025 → 10/21/2025
    - 21 October 2025 → 10/21/2025
    - 21 October, 2025 → 10/21/2025
    - Oct 21, 2025 → 10/21/2025
    - 2025-10-21 → 10/21/2025
    - 19th day of October, 2005 → 10/19/2005
    Also fixes common typos in month names (e.g., Seplember → September)
    """
    from datetime import datetime
    import re as date_re
    
    # Fix common month name typos (OCR/typing errors)
    month_typo_fixes = {
        'Seplember': 'September',
        'Septmber': 'September',
        'Sepetmber': 'September',
        'Janurary': 'January',
        'Febuary': 'February',
        'Feburary': 'February',
        'Apirl': 'April',
        'Augest': 'August',
        'Augst': 'August',
        'Ocotber': 'October',
        'Novmber': 'November',
        'Decmber': 'December',
    }
    for typo, correct in month_typo_fixes.items():
        date_str = date_re.sub(typo, correct, date_str, flags=date_re.IGNORECASE)
    
    # First, handle "Xth day of Month, Year" format
    day_of_match = date_re.search(r'(\d{1,2})(?:st|nd|rd|th)?\s+day\s+of\s+([A-Za-z]+),?\s+(\d{4})', date_str, date_re.IGNORECASE)
    if day_of_match:
        day = day_of_match.group(1)
        month = day_of_match.group(2)
        year = day_of_match.group(3)
        date_str = f"{month} {day}, {year}"
    
    # Try multiple date formats
    date_formats = [
        '%B %d, %Y',     # October 21, 2025
        '%B %d %Y',      # October 21 2025
        '%b %d, %Y',     # Oct 21, 2025
        '%b %d %Y',      # Oct 21 2025
        '%d %B, %Y',     # 21 October, 2025
        '%d %B %Y',      # 21 October 2025
        '%d %b, %Y',     # 21 Oct, 2025
        '%d %b %Y',      # 21 Oct 2025
        '%Y-%m-%d',      # 2025-10-21
        '%m/%d/%Y',      # 10/21/2025 (already correct)
        '%d/%m/%Y',      # 21/10/2025 (European)
    ]
    
    for fmt in date_formats:
        try:
            dt = datetime.strptime(date_str.strip(), fmt)
            return dt.strftime('%m/%d/%Y')
        except ValueError:
            continue
    
    # If no format matches, return original
    return date_str

def detect_gdpr_emea(text: str, address_blob: str = "") -> dict:
    """
    Detect contract jurisdiction region: EMEA, North America, or APAC.
    Also determines if GDPR applies (European/EMEA jurisdiction).
    """
    region_info = {}
    
    combined_text = f"{text} {address_blob}" if address_blob else text

    # Check for explicit region mentions
    if re.search(r'\bEMEA\b', combined_text, re.IGNORECASE):
        region_info['region'] = 'EMEA'
        region_info['subject_to_gdpr'] = True
        region_info['gdpr_reason'] = 'EMEA region'
    elif re.search(r'\bNorth America\b', combined_text, re.IGNORECASE):
        region_info['region'] = 'North America'
    elif re.search(r'\bAPAC\b|\bAsia Pacific\b', combined_text, re.IGNORECASE):
        region_info['region'] = 'APAC'
    
    # Check for configured primary EMEA legal entity trigger
    if PRIMARY_COMPANY_EMEA_PATTERN and re.search(PRIMARY_COMPANY_EMEA_PATTERN, combined_text, re.IGNORECASE):
        region_info['subject_to_gdpr'] = True
        region_info['gdpr_reason'] = PRIMARY_COMPANY_GDPR_REASON
        if 'region' not in region_info:
            region_info['region'] = 'EMEA'
    
    # Check for European legal entity types (avoid generic "Limited" matches)
    for pattern in EU_ENTITY_PATTERNS:
        match = re.search(pattern, combined_text)
        if match:
            if 'subject_to_gdpr' not in region_info:
                region_info['subject_to_gdpr'] = True
                entity_label = match.group(0).split()[-1].rstrip('.').upper()
                region_info['gdpr_reason'] = f'European legal entity type ({entity_label})'
            if 'region' not in region_info:
                region_info['region'] = 'EMEA'
            break
    
    # Check for European countries
    for country in EMEA_COUNTRIES:
        if re.search(rf'\b{country}\b', combined_text, re.IGNORECASE):
            if 'subject_to_gdpr' not in region_info:
                region_info['subject_to_gdpr'] = True
                region_info['gdpr_reason'] = f'European address ({country})'
            if 'region' not in region_info:
                region_info['region'] = 'EMEA'
            break
    
    # Check for North American states/countries
    for location in NA_LOCATIONS:
        if re.search(rf'\b{location}\b', combined_text, re.IGNORECASE):
            if 'region' not in region_info:
                region_info['region'] = 'North America'
            break
    
    # Check for APAC countries
    for country in APAC_COUNTRIES:
        if re.search(rf'\b{country}\b', combined_text, re.IGNORECASE):
            if 'region' not in region_info:
                region_info['region'] = 'APAC'
            break
    
    return region_info

def extract_term_info(text: str) -> dict:
    """
    ENHANCED: Extract ALL term-related information using comprehensive regex patterns.
    
    Captures contract term patterns like:
    - "initial term of five (5) years (the "Initial Term")"
    - "automatically renew for additional one (1) year terms (each, a "Renewal Term")"
    - "thirty (30) days prior written notice"
    - "until terminated" with notice periods
    - "shall continue for a period of 36 months"
    
    Returns: initial_term, renewal_term, renewal_notice_period, termination_notice_period,
             payment_terms, late_charge, termination_for_convenience, auto_renewal,
             contract_duration, perpetual_notice_period
    """
    term_info = {}
    
    # Normalize text for better matching
    text_normalized = re.sub(r'\s+', ' ', text)
    
    # ========================================================================
    # DURATION PATTERNS - Common ways to express time periods
    # ========================================================================
    # Matches: "thirty (30) days", "five (5) years", "36 months", "one year", "1 year"
    DURATION_WORD_NUM = r'(?:one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|' \
                        r'fifteen|eighteen|twenty|twenty-?four|thirty|thirty-?six|forty-?five|' \
                        r'sixty|ninety|one\s+hundred\s+(?:twenty|eighty))\s*\(\d+\)'
    DURATION_NUM_ONLY = r'\d{1,3}'
    DURATION_PERIOD = r'(?:days?|months?|years?)'
    # Full duration pattern: "thirty (30) days" or "5 years" or "36 months"
    DURATION_FULL = rf'(?:(?:{DURATION_WORD_NUM}|{DURATION_NUM_ONLY})\s*{DURATION_PERIOD})'
    def capture_snippet_before(marker: str, window: int = 300) -> Optional[str]:
        """Capture text before a defined term marker like ("Initial Term").
        
        Looks for the FIRST occurrence which is typically the definition.
        Excludes matches that are just references (like "the Initial Term" in later text).
        Handles markdown formatting like **Term** and various quote styles.
        """
        # Pattern for defined term: ... (the "Term") or ... (each, a "Term") or ... ("Term")
        # Also handles markdown: ... (each, a " **Term** ") or ... (" **Term** ")
        # Quote chars: " ' " " and optional markdown ** around the term
        quote_chars = r'["\u201C\u201D\'\s]*'
        markdown_bold = r'(?:\*\*)?'
        pattern = rf'([\s\S]{{1,{window}}})\s*\({quote_chars}(?:each\s*,?\s*a\s+|the\s+|a\s+)?{quote_chars}{markdown_bold}\s*{re.escape(marker)}\s*{markdown_bold}{quote_chars}\)'
        
        # Find ALL matches and choose the best one (typically the first definition-style match)
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        
        best_snippet = None
        for match in matches:
            snippet = match.group(1)
            snippet = re.sub(r'\s+', ' ', snippet)
            # Remove markdown formatting from snippet
            snippet = re.sub(r'\*\*', '', snippet)
            
            # Keep only the trailing clause after the last sentence delimiter
            for delimiter in '.;:':
                idx = snippet.rfind(delimiter)
                if idx != -1 and len(snippet) - idx <= 250:
                    snippet = snippet[idx + 1:]
                    break
            snippet = snippet.strip(' ,;')
            
            # Check if this is a DEFINITION (has duration language) vs just a REFERENCE
            # Definition patterns: "shall have", "will be", "shall be", "for a/an", "of X years"
            snippet_lower = snippet.lower()
            is_definition = any(kw in snippet_lower for kw in [
                'shall have', 'will have', 'shall be', 'will be',
                'term of', 'period of', 'for a period', 'for an additional',
                'renew for', 'automatically renew', 'initial term of'
            ])
            
            # Reference patterns: "the end of", "prior to the", "during the", "from the"
            is_reference = any(kw in snippet_lower for kw in [
                'end of the', 'from the end', 'prior to the', 'during the',
                'within the', 'throughout the', 'at the end'
            ])
            
            if is_definition and not is_reference:
                best_snippet = snippet
                break  # Use first definition match
            elif best_snippet is None and not is_reference:
                best_snippet = snippet  # Keep first non-reference as fallback
        
        return best_snippet
    
    def extract_duration_from_text(snippet: str) -> Optional[str]:
        """Extract a duration like 'five (5) years' from text"""
        if not snippet:
            return None
        # Pattern: word (number) period - e.g., "five (5) years"
        match = re.search(rf'({DURATION_WORD_NUM}\s*{DURATION_PERIOD})', snippet, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        # Pattern: just number period - e.g., "36 months" or "5 years"
        match = re.search(rf'(\d{{1,3}}\s*{DURATION_PERIOD})', snippet, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        return None

    # ========================================================================
    # INITIAL TERM EXTRACTION
    # ========================================================================
    # Get defined terms for Initial Term and Renewal Term
    defined_terms = extract_defined_terms(text)
    
    # Try to extract from defined term ("Initial Term")
    if 'Initial Term' not in defined_terms:
        fallback_initial = capture_snippet_before('Initial Term')
        if fallback_initial:
            defined_terms['Initial Term'] = fallback_initial

    if 'Initial Term' in defined_terms:
        duration = extract_duration_from_text(defined_terms['Initial Term'])
        if duration:
            term_info['initial_term'] = duration
        else:
            term_info['initial_term'] = defined_terms['Initial Term']
    
    # Fallback patterns for initial term if not found via defined terms
    if 'initial_term' not in term_info:
        # PRIORITY 1: NDA/CDA specific patterns - "shall expire X years after Effective Date"
        # These are HIGH confidence and should be checked first
        nda_term_patterns = [
            # "This Agreement shall expire two (2) years after the Effective Date"
            rf'(?:this\s+)?agreement\s+shall\s+expire\s+({DURATION_FULL})\s+after\s+(?:the\s+)?effective\s+date',
            # "shall be effective for a period of two (2) years from the Effective Date"
            rf'shall\s+be\s+effective\s+for\s+(?:a\s+)?(?:period|term)\s+of\s+({DURATION_FULL})\s+from\s+(?:the\s+)?effective\s+date',
            # "shall remain in effect for two (2) years from the Effective Date"
            rf'shall\s+remain\s+in\s+(?:full\s+)?(?:force\s+and\s+)?effect\s+for\s+({DURATION_FULL})\s+from\s+(?:the\s+)?effective\s+date',
        ]
        for pattern in nda_term_patterns:
            match = re.search(pattern, text_normalized, re.IGNORECASE)
            if match:
                term_info['initial_term'] = match.group(1).strip()
                break
        
        # PRIORITY 2: Standard initial term patterns (skip if contains survival language)
        if 'initial_term' not in term_info:
            initial_term_patterns = [
                # "initial term of five (5) years"
                rf'initial\s+term\s+of\s+({DURATION_FULL})',
                # "for an initial period of 36 months"
                rf'initial\s+(?:term|period)\s+of\s+({DURATION_FULL})',
                # "Term of this Agreement shall be 5 years"
                rf'term\s+(?:of\s+)?(?:this\s+)?(?:agreement|contract)\s+(?:shall\s+be|is|shall\s+commence\s+.*?and\s+continue\s+for)\s+({DURATION_FULL})',
                # "effective for a period of one (1) year"
                rf'effective\s+for\s+(?:a\s+)?(?:period|term)\s+of\s+({DURATION_FULL})',
                # "shall have an initial term of one (1) year"
                rf'shall\s+have\s+an?\s+initial\s+term\s+of\s+({DURATION_FULL})',
            ]
            for pattern in initial_term_patterns:
                match = re.search(pattern, text_normalized, re.IGNORECASE)
                if match:
                    # Check context to exclude survival/confidentiality clauses
                    context_start = max(0, match.start() - 100)
                    context_end = min(len(text_normalized), match.end() + 100)
                    context = text_normalized[context_start:context_end].lower()
                    
                    # EXCLUDE if this is a survival/confidentiality clause
                    survival_keywords = ['surviv', 'after expiration', 'after termination', 
                                        'continue for a period of', 'confidential']
                    is_survival_clause = any(kw in context for kw in survival_keywords)
                    
                    if not is_survival_clause:
                        term_info['initial_term'] = match.group(1).strip()
                        break
        
        # PRIORITY 3: General "shall continue for" pattern (with survival exclusion)
        if 'initial_term' not in term_info:
            continue_pattern = rf'(?:shall\s+)?continue\s+for\s+(?:a\s+)?(?:period|term)\s+of\s+({DURATION_FULL})'
            for match in re.finditer(continue_pattern, text_normalized, re.IGNORECASE):
                # Check context to exclude survival clauses
                context_start = max(0, match.start() - 150)
                context_end = min(len(text_normalized), match.end() + 100)
                context = text_normalized[context_start:context_end].lower()
                
                # EXCLUDE if this is about survival/confidentiality
                if not any(kw in context for kw in ['surviv', 'after expiration', 'after termination', 
                                                     'obligations shall', 'confidential']):
                    term_info['initial_term'] = match.group(1).strip()
                    break
    
    # ========================================================================
    # RENEWAL TERM EXTRACTION
    # ========================================================================
    if 'Renewal Term' not in defined_terms:
        fallback_renewal = capture_snippet_before('Renewal Term')
        if fallback_renewal:
            defined_terms['Renewal Term'] = fallback_renewal

    if 'Renewal Term' in defined_terms:
        duration = extract_duration_from_text(defined_terms['Renewal Term'])
        if duration:
            term_info['renewal_term'] = duration
        else:
            term_info['renewal_term'] = defined_terms['Renewal Term']
    
    # Fallback patterns for renewal term
    if 'renewal_term' not in term_info:
        renewal_term_patterns = [
            # "automatically renew for additional one (1) year terms" - must have "term" or "period" after duration
            rf'(?:automatically\s+)?renew(?:s|ed)?\s+for\s+(?:additional\s+|successive\s+)?({DURATION_FULL})\s+(?:terms?|periods?)',
            # "renewed for successive one year periods"
            rf'renew(?:s|ed)?\s+for\s+(?:successive|additional)\s+({DURATION_FULL})\s+(?:terms?|periods?)',
            # "each renewal term shall be one (1) year"
            rf'(?:each\s+)?renewal\s+term\s+(?:shall\s+be|of)\s+({DURATION_FULL})',
            # "renew for an additional one (1) year term"
            rf'renew\s+for\s+an?\s+additional\s+({DURATION_FULL})\s+term',
        ]
        for pattern in renewal_term_patterns:
            match = re.search(pattern, text_normalized, re.IGNORECASE)
            if match:
                duration = match.group(1).strip()
                # Validate it's not just a notice period (days are rarely renewal terms)
                if 'day' not in duration.lower() or 'year' in duration.lower() or 'month' in duration.lower():
                    term_info['renewal_term'] = duration
                    break
                # Check if it has "term" or "period" qualifier to confirm it's renewal term
                full_match = match.group(0).lower()
                if 'term' in full_match or 'period' in full_match:
                    term_info['renewal_term'] = duration
                    break
    
    # ========================================================================
    # NOTICE PERIOD EXTRACTION - COMPREHENSIVE
    # ========================================================================
    # Common notice patterns: "thirty (30) days", "sixty (60) days", "ninety (90) days"
    
    NOTICE_DURATION = rf'(?:(?:{DURATION_WORD_NUM}|{DURATION_NUM_ONLY})\s*(?:days?|months?))'
    
    # Pattern 1: Renewal notice - "ninety (90) days' prior written notice of non-renewal"
    renewal_notice_patterns = [
        rf"({NOTICE_DURATION})['\u2019]?\s+(?:prior\s+)?(?:written\s+)?notice\s+(?:of\s+)?(?:non[- ]?renewal|intent\s+not\s+to\s+renew)",
        rf'notice\s+of\s+non[- ]?renewal[^.]*?({NOTICE_DURATION})',
        rf'(?:non[- ]?renewal|not\s+to\s+renew)[^.]*?({NOTICE_DURATION})\s+(?:prior\s+)?(?:written\s+)?notice',
    ]
    for pattern in renewal_notice_patterns:
        match = re.search(pattern, text_normalized, re.IGNORECASE)
        if match:
            term_info['renewal_notice_period'] = match.group(1).strip()
            break
    
    # Pattern 2: Termination notice - comprehensive patterns
    termination_notice_patterns = [
        # "terminate this Agreement without cause with sixty (60) days prior written notice"
        rf'terminat(?:e|ion)[^.]*?(?:without\s+cause\s+)?(?:with|upon|by\s+giving)\s+({NOTICE_DURATION})[\'s]?\s+(?:prior\s+)?(?:written\s+)?notice',
        # "sixty (60) days prior written notice of termination"
        rf'({NOTICE_DURATION})[\'s]?\s+(?:prior\s+)?(?:written\s+)?notice\s+(?:of\s+)?terminat(?:e|ion)',
        # "upon thirty (30) days' notice"
        rf'upon\s+({NOTICE_DURATION})[\'s]?\s+(?:prior\s+)?(?:written\s+)?notice',
        # "by giving at least 30 days notice"
        rf'(?:by\s+)?giving\s+(?:at\s+least\s+)?({NOTICE_DURATION})[\'s]?\s+(?:prior\s+)?(?:written\s+)?notice',
        # "with 2 (two) months' notice"
        rf'with\s+({DURATION_NUM_ONLY}\s*\(\w+\)\s*(?:days?|months?))[\'s]?\s+notice',
        # "effective thirty (30) days from receipt"
        rf'(?:terminat(?:e|ion)\s+)?(?:shall\s+be\s+)?effective\s+({NOTICE_DURATION})\s+(?:from|after)\s+(?:receipt|notice)',
        # General: "at least 30 days prior notice"
        rf'at\s+least\s+({NOTICE_DURATION})[\'s]?\s+(?:prior\s+)?(?:written\s+)?notice',
    ]
    if 'termination_notice_period' not in term_info:
        for pattern in termination_notice_patterns:
            match = re.search(pattern, text_normalized, re.IGNORECASE)
            if match:
                term_info['termination_notice_period'] = match.group(1).strip()
                break
    
    # ========================================================================
    # PERPETUAL/INDEFINITE CONTRACT DETECTION
    # ========================================================================
    perpetual_patterns = [
        # "shall continue until terminated by either party"
        rf'(?:shall\s+)?continue\s+(?:in\s+(?:full\s+)?(?:force|effect)\s+)?until\s+terminat(?:e|ed|ion)',
        # "continue indefinitely until terminated"
        r'continue\s+indefinitely',
        # "perpetual license"
        r'perpetual\s+(?:license|agreement|term)',
        # "remain in effect until terminated"
        r'remain\s+in\s+(?:full\s+)?(?:force\s+and\s+)?effect\s+until\s+terminat',
        # "no fixed term"
        r'no\s+fixed\s+(?:term|expiration)',
        # "until terminated" (standalone)
        r'(?:shall\s+|will\s+)?(?:be\s+|remain\s+)?(?:effective\s+|valid\s+|in\s+force\s+)?until\s+(?:either\s+party\s+)?terminat(?:e|ed|es|ion)',
        # "until either party terminates"
        r'until\s+(?:such\s+time\s+as\s+)?(?:either\s+)?party\s+(?:gives\s+notice|terminates)',
        # "effective until" patterns  
        r'effective\s+until\s+terminat',
        # "in force until"
        r'in\s+force\s+until\s+terminat',
        # NEW: "may be terminated upon X days notice" without fixed end date (evergreen)
        r'(?:agreement|contract)\s+may\s+be\s+terminated\s+upon\s+.*?notice\s+by\s+either\s+party',
        # NEW: "terminated upon written notice" pattern
        r'(?:may|can)\s+be\s+terminated\s+(?:at\s+any\s+time\s+)?(?:upon|with|by\s+giving)\s+.*?(?:days?|months?)\s+.*?notice',
        # NEW: Agreement continues until termination notice (no fixed term)
        r'(?:this\s+)?agreement\s+(?:shall\s+)?(?:continue|remain)\s+(?:in\s+effect\s+)?(?:unless|until)\s+(?:and\s+until\s+)?terminat',
        # NEW: Evergreen with notice
        r'evergreen\s+(?:agreement|contract|term)',
    ]
    
    for pattern in perpetual_patterns:
        match = re.search(pattern, text_normalized, re.IGNORECASE)
        if match:
            term_info['perpetual'] = True
            # Look for associated notice period
            context = text_normalized[max(0, match.start()-50):min(len(text_normalized), match.end()+300)]
            notice_match = re.search(rf'({NOTICE_DURATION})[\'s]?\s+(?:prior\s+)?(?:written\s+)?notice', context, re.IGNORECASE)
            if notice_match:
                term_info['perpetual_notice_period'] = notice_match.group(1).strip()
            break
    
    # ========================================================================
    # AUTO-RENEWAL DETECTION
    # ========================================================================
    auto_renew_patterns = [
        r'(?:will\s+|shall\s+)?automatically\s+renew',
        r'auto[- ]?renew(?:s|al)?',
        r'renew(?:s|ed)?\s+automatically',
        r'successive\s+(?:renewal\s+)?(?:terms?|periods?)',
    ]
    
    for pattern in auto_renew_patterns:
        match = re.search(pattern, text_normalized, re.IGNORECASE)
        if match:
            term_info['auto_renewal'] = True
            # Try to extract renewal term from context
            context = text_normalized[match.start():min(len(text_normalized), match.end()+200)]
            duration_match = re.search(rf'(?:additional\s+|successive\s+)?({DURATION_FULL})\s*(?:terms?|periods?)?', context, re.IGNORECASE)
            if duration_match and 'renewal_term' not in term_info:
                term_info['renewal_term'] = duration_match.group(1).strip()
            break
    
    # ========================================================================
    # EFFECTIVE DATE EXTRACTION (from preamble)
    # ========================================================================
    effective_date_patterns = [
        # "effective as of October 19, 2005 (the "Effective Date")"
        r'effective\s+(?:as\s+of\s+)?([A-Za-z]+\s+\d{1,2},?\s+\d{4})\s*\((?:the\s+)?["\']?Effective\s+Date',
        # "dated October 19, 2005" or "dated as of October 19, 2005"
        r'(?:dated|entered\s+into)\s+(?:as\s+of\s+)?([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
        # "this 19th day of October, 2005"
        r'this\s+(\d{1,2}(?:st|nd|rd|th)?\s+day\s+of\s+[A-Za-z]+,?\s+\d{4})',
        # "Effective Date" shall mean [date]
        r'"?Effective\s+Date"?\s+(?:shall\s+)?mean[^.]*?([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
        # European format: "25 November 2024"
        r'(?:dated|effective|as\s+of)\s+(\d{1,2}\s+[A-Za-z]+\s+\d{4})',
    ]
    
    for pattern in effective_date_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            # Normalize to MM/DD/YYYY format
            term_info['effective_date_extracted'] = normalize_date_format(match.group(1).strip())
            break
    
    # ========================================================================
    # EXISTING PATTERNS (Payment terms, late charge, termination for convenience, governing law)
    # ========================================================================
    
    # Payment terms: "Payment shall be made in full thirty (30) days"
    payment_pattern = rf'payments?[^.]*?({NOTICE_DURATION})'
    payment_match = re.search(payment_pattern, text_normalized, re.IGNORECASE)
    if payment_match:
        term_info['payment_terms'] = payment_match.group(1)
    
    # Late charge
    late_charge_pattern = r'late\s+(?:charge|fee|payment)[^.]{0,100}'
    late_charge_match = re.search(late_charge_pattern, text_normalized, re.IGNORECASE)
    if late_charge_match:
        term_info['late_charge'] = late_charge_match.group(0).strip()
    
    # Termination for convenience
    termination_convenience_pattern = r'terminat(?:e|ing)\s+this\s+(?:Agreement|Contract)\s+(?:for\s+(?:any\s+reason|no\s+reason|convenience)|without\s+cause)[^.]{0,200}'
    termination_convenience_match = re.search(termination_convenience_pattern, text_normalized, re.IGNORECASE)
    term_info['termination_for_convenience'] = 'Yes' if termination_convenience_match else 'No'
    if termination_convenience_match:
        clause_text = termination_convenience_match.group(0).strip()
        notice_in_clause = re.search(rf'({NOTICE_DURATION})', clause_text, re.IGNORECASE)
        if notice_in_clause:
            term_info['termination_for_convenience_notice'] = notice_in_clause.group(1)
        term_info['termination_for_convenience_clause'] = clause_text[:250]
    
    # Governing law
    governing_law_pattern = r'govern(?:ed|ing)\s+(?:by|under)\s+(?:the\s+)?laws?\s+of\s+(?:the\s+)?(?:State\s+of\s+)?([A-Z][A-Za-z\s]+?)(?:\s+without|\s+excluding|\.|,|;|\s+and)'
    governing_law_match = re.search(governing_law_pattern, text, re.IGNORECASE)
    if governing_law_match:
        term_info['governing_law'] = governing_law_match.group(1).strip()
    
    # Purpose clause
    purpose_pattern = r'Purpose[)\]\.]\s+(?:In order to facilitate|To facilitate|For the purpose of)?.*?relat(?:ing|es?)\s+to\s+([^(]{10,200})\s*\((?:the\s+)?"Purpose"\)'
    purpose_match = re.search(purpose_pattern, text, re.IGNORECASE | re.DOTALL)
    if purpose_match:
        purpose_text = purpose_match.group(1).strip()
        purpose_text = re.sub(r'\s+', ' ', purpose_text)
        term_info['purpose'] = purpose_text
    
    # ========================================================================
    # SOW/PROJECT DATE EXTRACTION - Start Date and End Date
    # ========================================================================
    # Common SOW patterns: "Project start date shall be 25 November 2024 and shall end 28 February 2025"
    
    # Date pattern that handles: "25 November 2024", "25 November, 2024", "November 25, 2024", "11/25/2024"
    DATE_PATTERN = r'(?:\d{1,2}\s+[A-Za-z]+,?\s+\d{4}|[A-Za-z]+\s+\d{1,2},?\s+\d{4}|\d{1,2}/\d{1,2}/\d{4})'
    
    # Pattern 1: "Project start date shall be X and shall end Y"
    sow_dates_pattern = rf'(?:project\s+)?start\s+date\s+(?:shall\s+be|is)\s+({DATE_PATTERN})\s+and\s+(?:shall\s+)?end\s+({DATE_PATTERN})'
    sow_dates_match = re.search(sow_dates_pattern, text_normalized, re.IGNORECASE)
    if sow_dates_match:
        term_info['sow_start_date'] = normalize_date_format(sow_dates_match.group(1).strip())
        term_info['sow_end_date'] = normalize_date_format(sow_dates_match.group(2).strip())
    else:
        # Pattern 2: Separate start and end dates
        start_patterns = [
            rf'(?:project\s+)?start\s+date\s+(?:shall\s+be|is|:)\s*({DATE_PATTERN})',
            rf'(?:commencement|commence)\s+(?:date\s+)?(?:shall\s+be|is|on|:)\s*({DATE_PATTERN})',
            rf'services?\s+(?:shall\s+)?(?:commence|begin)\s+(?:on\s+)?({DATE_PATTERN})',
        ]
        for pattern in start_patterns:
            match = re.search(pattern, text_normalized, re.IGNORECASE)
            if match:
                term_info['sow_start_date'] = normalize_date_format(match.group(1).strip())
                break
        
        end_patterns = [
            rf'(?:project\s+)?end\s+date\s+(?:shall\s+be|is|:)\s*({DATE_PATTERN})',
            rf'(?:shall\s+)?end\s+(?:on\s+)?({DATE_PATTERN})',
            rf'completion\s+date\s+(?:shall\s+be|is|of|:)\s*({DATE_PATTERN})',
            rf'services?\s+(?:shall\s+)?(?:end|terminate|conclude)\s+(?:on\s+)?({DATE_PATTERN})',
        ]
        for pattern in end_patterns:
            match = re.search(pattern, text_normalized, re.IGNORECASE)
            if match:
                term_info['sow_end_date'] = normalize_date_format(match.group(1).strip())
                break
    
    # Pattern 3: "will terminate upon completion" - mark as completion-based
    completion_patterns = [
        r'(?:shall|will)\s+terminate\s+upon\s+completion',
        r'terminat(?:e|es)\s+upon\s+(?:the\s+)?completion\s+of\s+(?:the\s+)?services',
        r'until\s+(?:the\s+)?completion\s+of\s+(?:the\s+)?(?:project|services|work)',
    ]
    for pattern in completion_patterns:
        if re.search(pattern, text_normalized, re.IGNORECASE):
            term_info['sow_completion_based'] = True
            break
    
    # ========================================================================
    # PARENT DOCUMENT REFERENCE EXTRACTION (for SOWs/Amendments)
    # ========================================================================
    # SOWs and Amendments reference their governing Master Agreement
    # Pattern: "subject to the terms of the Master Services Agreement dated October 16, 2015"
    
    parent_ref_patterns = [
        # "Master Services Agreement... dated [date]"
        rf'(?:Master\s+)?(?:Services?\s+)?Agreement.*?(?:dated|effective)\s+(?:as\s+of\s+)?({DATE_PATTERN})',
        # "pursuant to the Agreement dated [date]"
        rf'pursuant\s+to\s+(?:the\s+)?(?:Master\s+)?(?:Services?\s+)?Agreement.*?dated\s+({DATE_PATTERN})',
        # "under the MSA dated [date]"
        rf'under\s+(?:the\s+)?(?:MSA|Master\s+(?:Services?\s+)?Agreement).*?dated\s+({DATE_PATTERN})',
        # "to which this Statement of Work is appended" - look for date nearby
        rf'Statement\s+of\s+Work.*?(?:appended|attached|subject).*?(?:dated|effective)\s+(?:as\s+of\s+)?({DATE_PATTERN})',
    ]
    
    for pattern in parent_ref_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            parent_date = normalize_date_format(match.group(1).strip())
            term_info['parent_agreement_date'] = parent_date
            
            # Determine parent agreement type
            full_match = match.group(0).lower()
            if 'master' in full_match or 'msa' in full_match:
                term_info['parent_agreement_type'] = 'MSA'
            elif 'service' in full_match:
                term_info['parent_agreement_type'] = 'Services Agreement'
            else:
                term_info['parent_agreement_type'] = 'Agreement'
            break
    
    return term_info


def extract_defined_terms(text: str) -> dict:
    """Extract defined terms of the form ... (the "Term")."""
    defined_terms: Dict[str, str] = {}
    term_pattern = r'\((?:the\s+)?["\u201C\u201D]?([A-Za-z][A-Za-z\s&\-]+)["\u201C\u201D]?\)'
    party_context = get_party_clauses(text)

    for match in re.finditer(term_pattern, text):
        term_name = match.group(1).strip()
        if not term_name:
            continue

        window_start = max(0, match.start() - 400)
        snippet_context = text[window_start:match.start()]
        lines = [line.strip() for line in snippet_context.splitlines() if line.strip()]
        if lines:
            snippet = ' '.join(lines[-3:])
        else:
            snippet = snippet_context.strip()
        snippet = re.sub(r'\s+', ' ', snippet).strip()
        snippet = re.sub(r'^(?:and|or)\s+', '', snippet, flags=re.IGNORECASE)

        lower_term = term_name.lower()
        if lower_term in {'effective date', 'expiration date', 'termination date'}:
            date_literal = _extract_date_literal(snippet)
            if not date_literal:
                lookahead = text[match.end():match.end() + 200]
                date_literal = _extract_date_literal(lookahead)
            if date_literal:
                snippet = date_literal
        elif lower_term in {'initial term', 'renewal term'}:
            term_match = re.search(r'((?:one|two|three|four|five|six|seven|eight|nine|ten|\d+)\s*\(\d+\)\s+\w+(?:\s+\w+)*)', snippet, re.IGNORECASE)
            if term_match:
                snippet = term_match.group(1)

        if lower_term in PARTY_ROLE_TERMS_LOWER:
            if lower_term in PRIMARY_COMPANY_ROLE_HINTS and party_context.get('primary_party'):
                snippet = party_context['primary_party']
            elif lower_term in COUNTERPARTY_ROLE_HINTS and party_context.get('counterparty'):
                snippet = party_context['counterparty']
            snippet = re.sub(r',\s+a\s+\w+\s+(?:corporation|company|llc|limited|partnership|entity).*$', '', snippet, flags=re.IGNORECASE)
            snippet = re.sub(r',\s+having.*$', '', snippet, flags=re.IGNORECASE)
            snippet = re.sub(r',\s+with.*$', '', snippet, flags=re.IGNORECASE)

        snippet = snippet.strip(' ,;')
        if snippet:
            defined_terms[term_name] = snippet

    return defined_terms


def extract_parties(text: str) -> str:
    """Return normalized party names for metadata display."""
    party_context = get_party_clauses(text)
    party_details = _build_party_details_from_context(party_context)
    defined_terms = extract_defined_terms(text)

    parties: List[str] = []

    for key in ('party1_name', 'party2_name', 'primary_party_name', 'counterparty_name'):
        candidate = _normalize_party_candidate(party_details.get(key))
        if candidate and candidate not in parties:
            parties.append(candidate)
        if len(parties) >= 2:
            break

    if len(parties) < 2:
        ordered_info = party_context.get('ordered_info') or []
        for info in ordered_info:
            candidate = _normalize_party_candidate(info.get('name'))
            if candidate and candidate not in parties:
                parties.append(candidate)
            if len(parties) >= 2:
                break

    if len(parties) < 2:
        for role in PARTY_ROLE_TERMS:
            candidate = _normalize_party_candidate(defined_terms.get(role))
            if candidate and candidate not in parties:
                parties.append(candidate)
            if len(parties) >= 2:
                break

    if len(parties) < 2:
        for term_value in defined_terms.values():
            candidate = _normalize_party_candidate(term_value)
            if candidate and candidate not in parties:
                parties.append(candidate)
            if len(parties) >= 2:
                break

    if len(parties) < 2:
        between_pattern = r'between\s+(.{1,600}?)\s+(?:and|&)\s+(.{1,600}?)(?:\.|,|\n)'
        between_match = re.search(between_pattern, text, re.IGNORECASE | re.DOTALL)
        if between_match:
            for candidate in between_match.groups():
                normalized = _normalize_party_candidate(candidate)
                if normalized and normalized not in parties:
                    parties.append(normalized)
                if len(parties) >= 2:
                    break

    if len(parties) < 2:
        multiline_pattern = (
            r'This\s+[A-Za-z\s]*?Agreement[\s\S]{0,4000}?(?:between|by\s+and\s+between)\s+'
            r'(.{30,2000}?)\s+(?:and|&)\s+(.{30,2000}?)(?:\.|\n{2}|;|\r)'
        )
        for match in re.finditer(multiline_pattern, text, re.IGNORECASE | re.DOTALL):
            for candidate in match.groups():
                normalized = _normalize_party_candidate(candidate)
                if normalized and normalized not in parties:
                    parties.append(normalized)
            if len(parties) >= 2:
                break

    if len(parties) < 2 and PRIMARY_COMPANY_NAMES:
        for name in PRIMARY_COMPANY_NAMES:
            pattern = re.compile(rf'([A-Z0-9][A-Za-z0-9&.,\s-]{{0,120}}{re.escape(name)}[A-Za-z0-9&.,\s-]{{0,40}})', re.IGNORECASE)
            match = pattern.search(text)
            if match:
                normalized = _normalize_party_candidate(match.group(0))
                if normalized and normalized not in parties:
                    parties.append(normalized)
            if len(parties) >= 2:
                break

    if len(parties) < 2:
        legal_entity_pattern = r'([A-Z][A-Za-z0-9&.,\s]+(?:Corporation|Corp\.?|Company|LLC|Limited|Ltd\.?|GmbH|ApS|AG|BV|Inc\.?|SAS|SA))'
        for match in re.finditer(legal_entity_pattern, text, re.IGNORECASE):
            normalized = _normalize_party_candidate(match.group(1))
            if normalized and normalized not in parties:
                parties.append(normalized)
            if len(parties) >= 2:
                break

    parties = _dedupe_candidates(parties)
    parties = _prioritize_primary_company(parties)

    return " AND ".join(parties[:2]) if parties else "Unknown"


def _build_party_details_from_context(context: Dict[str, Any]) -> Dict[str, str]:
    details: Dict[str, str] = {}

    ordered = context.get('ordered_info') or []
    if ordered:
        primary = ordered[0]
        if primary.get('name'):
            details['party1_name'] = primary['name']
        if primary.get('address'):
            details['party1_address'] = primary['address']
    if len(ordered) > 1:
        secondary = ordered[1]
        if secondary.get('name'):
            details['party2_name'] = secondary['name']
        if secondary.get('address'):
            details['party2_address'] = secondary['address']

    primary_info = context.get('primary_party_info') or {}
    counterparty_info = context.get('counterparty_info') or {}
    _add_primary_party_field(details, 'name', primary_info.get('name'))
    _add_primary_party_field(details, 'address', primary_info.get('address'))
    if counterparty_info.get('name'):
        details['counterparty_name'] = counterparty_info['name']
    if counterparty_info.get('address'):
        details['counterparty_address'] = counterparty_info['address']

    return details


def extract_party_details(text: str) -> Dict[str, str]:
    """Return normalized party name/address pairs."""
    context = get_party_clauses(text)
    return _build_party_details_from_context(context)


# ============================================================================
# TOOLS WITH RAG ENHANCEMENT
# ============================================================================

# Global vector DB instance
vector_db: Optional[ContractVectorDB] = None

def upload_contract(
    file_path: str,
    use_ocr: bool = False,
    tool_context: Optional[ToolContext] = None,
    relative_label: Optional[str] = None,
    force_reprocess: bool = False
) -> str:
    """
    Upload a contract (PDF or DOCX) to the system with AUTO-DETECTION of metadata.
    Extracts text, tables, auto-detects parties and type, generates embeddings, stores in ChromaDB.
    
    Args:
        file_path: Path to PDF or DOCX file
        use_ocr: Use OCR for scanned PDFs
        relative_label: Optional display label (e.g., relative path) for metadata
    
    Returns:
        Status message with contract ID and detected metadata
    """
    global vector_db
    
    if not vector_db:
        return "❌ Vector database not initialized"
    
    file_name = Path(file_path).name
    if file_name.startswith('~$') or file_name.startswith('.~$'):
        return f"ℹ️  Skipped temporary Office lock file: {file_name}"

    file_ext = Path(file_path).suffix.lower()
    
    # Extract based on file type
    if file_ext == '.pdf':
        if not PDF_SUPPORT:
            return "❌ PDF support not available. Install PyMuPDF: pip install pymupdf"
        doc_data = extract_text_from_pdf(file_path, use_ocr=use_ocr)
        text = doc_data["text"]
        tables = doc_data["tables"]
        metadata_extra = {
            "page_count": doc_data["page_count"],
            "has_tables": len(doc_data["tables"]) > 0,
            "table_count": len(doc_data["tables"]),
            "ocr_applied": doc_data.get("ocr_applied", False)
        }
        pdf_meta = doc_data.get("metadata") or {}
        for key, value in pdf_meta.items():
            if value:
                safe_key = f"pdf_{key.lower()}" if isinstance(key, str) else f"pdf_meta_{key}"
                metadata_extra[safe_key] = value
    elif file_ext in ['.docx', '.doc']:
        if not DOCX_SUPPORT:
            return "❌ Word support not available. Install python-docx: pip install python-docx"
        doc_data = extract_text_from_docx(file_path)
        text = doc_data["text"]
        tables = doc_data["tables"]
        metadata_extra = {
            "has_tables": len(doc_data["tables"]) > 0,
            "table_count": len(doc_data["tables"]),
            "doc_author": doc_data["metadata"].get("author", ""),
            "is_draft": True,  # Assume Word docs are drafts
            "ocr_applied": False
        }
        doc_meta = doc_data.get("metadata") or {}
        if doc_meta.get("created"):
            metadata_extra["doc_created"] = doc_meta["created"]
        if doc_meta.get("modified"):
            metadata_extra["doc_modified"] = doc_meta["modified"]
    else:
        return f"❌ Unsupported file type: {file_ext}. Use .pdf or .docx"
    
    if not text or len(text.strip()) < 50:
        return f"❌ Could not extract meaningful text from {file_path}"
    
    # Get folder name for type detection (agreement_folder is more accurate than OCR)
    agreement_folder = Path(file_path).parent.name
    
    # PHASE 1: Quick regex extraction (fast, deterministic)
    print("🔍 PHASE 1: Quick metadata extraction (regex-based)...")
    detected_type = detect_contract_type(text, folder_name=agreement_folder)
    detected_parties = extract_parties(text)
    detected_party_details = extract_party_details(text)
    detected_dates = extract_dates(text)
    detected_terms = extract_term_info(text)
    address_blob = " ".join(
        [value for key, value in detected_party_details.items() if key.endswith('address') and value]
    ) if detected_party_details else ""
    gdpr_info = detect_gdpr_emea(text, address_blob)
    sharepoint_fields = derive_sharepoint_metadata(detected_type, detected_party_details, gdpr_info)
    doc_context = derive_document_context(file_path, text, relative_label)
    
    print(f"   ✅ Type: {detected_type}")
    print(f"   ✅ Parties: {detected_parties}")
    if detected_dates:
        date_pairs = [f"{_humanize_key(key)}: {value}" for key, value in detected_dates.items()]
        print(f"   ✅ Dates: {', '.join(date_pairs)}")
    if detected_terms:
        term_pairs: List[str] = []
        for key, value in detected_terms.items():
            if isinstance(value, bool):
                display_value = 'Yes' if value else 'No'
            else:
                display_value = str(value)
            term_pairs.append(f"{_humanize_key(key)}: {display_value}")
        print(f"   ✅ Terms: {', '.join(term_pairs)}")
    if detected_terms.get('contract_term') == 'Perpetual':
        print(f"   ♾️  Contract: Perpetual (notice: {detected_terms.get('perpetual_notice_period')})")
    if detected_terms.get('governing_law'):
        print(f"   ⚖️  Governing Law: {detected_terms.get('governing_law')}")
    if detected_terms.get('purpose'):
        print(f"   🎯 Purpose: {detected_terms.get('purpose')[:50]}...")
    effective_region = sharepoint_fields.get('region') or gdpr_info.get('region')
    if effective_region:
        print(f"   🌍 Region: {effective_region}")
    if gdpr_info.get('subject_to_gdpr'):
        print(f"   ⚠️  GDPR: {gdpr_info.get('gdpr_reason')}")
    elif sharepoint_fields.get('gdpr_applicable') == 'Yes':
        print(f"   ⚠️  GDPR: {sharepoint_fields.get('gdpr_reason', 'EMEA address detected')}")
    if doc_context.get('document_status_label'):
        print(f"   📂 Status: {doc_context['document_status_label']}")
    if doc_context.get('document_role'):
        print(f"   📎 Document Role: {doc_context['document_role']}")
    if doc_context.get('related_contract_hint'):
        print(f"   🔗 Relationship: {doc_context['related_contract_hint']}")
    if doc_context.get('document_sequence_label'):
        sequence_details = doc_context['document_sequence_label']
        if doc_context.get('document_sequence_token') and doc_context['document_sequence_token'] not in sequence_details:
            sequence_details += f" (token: {doc_context['document_sequence_token']})"
        print(f"   🔢 Sequence: {sequence_details}")
    owner_display = sharepoint_fields.get('sharepoint_owners') or sharepoint_fields.get('sharepoint_owner_primary')
    if owner_display:
        dept = sharepoint_fields.get('sharepoint_department', 'Legal')
        print(f"   🗂️  SharePoint Owners: {owner_display} [{dept}]")
    if detected_terms.get('termination_for_convenience'):
        tfc = detected_terms['termination_for_convenience']
        notice = detected_terms.get('termination_for_convenience_notice')
        suffix = f" (notice: {notice})" if notice else ''
        print(f"   ✂️  Termination for Convenience: {tfc}{suffix}")
    
    # Generate deterministic contract ID
    source_hash = _compute_source_hash(file_path, text)
    contract_id = _generate_contract_id(file_path, text, source_hash)

    if vector_db.contract_exists(contract_id):
        if not force_reprocess:
            return f"ℹ️  Contract already ingested as {contract_id}. Skipping duplicate."
        print(f"♻️  Force reprocess enabled; replacing existing record {contract_id}.")
        vector_db.delete_contract(contract_id)
    
    # Build comprehensive metadata with PHASE 1 regex extractions
    metadata = {
        "contract_type": detected_type,
        "parties": detected_parties,
        "upload_date": datetime.now().isoformat(),
        "file_path": file_path,
        "filename": Path(file_path).name,
        "file_type": file_ext,
        "char_count": len(text),
        "extraction_phase": "regex",  # Phase 1 complete, awaiting Phase 2 AI validation
        "source_hash": source_hash,
        "relative_label": relative_label,
        **metadata_extra
    }
    
    # Add extracted dates, terms, and GDPR info if found
    if detected_dates:
        metadata.update(detected_dates)
    if detected_terms:
        metadata.update(detected_terms)
    if gdpr_info:
        metadata.update(gdpr_info)
    if detected_party_details:
        metadata.update(detected_party_details)
    if sharepoint_fields:
        metadata.update(sharepoint_fields)
    if doc_context:
        metadata.update({k: v for k, v in doc_context.items() if v not in (None, '', [])})
    
    if metadata_extra.get("ocr_applied"):
        metadata["ocr_applied"] = True

    # ENHANCEMENT: Extract dates from folder/file name if not found in text
    # Common pattern: "Agr Oct 19, 2005 - TERMINATED" or "2023-Q1 Renewal"
    if not detected_dates or not detected_dates.get('effective_date'):
        folder_path = Path(file_path).parent.name  # Get immediate parent folder
        filename_stem = Path(file_path).stem  # Get filename without extension
        combined_path_text = f"{folder_path} {filename_stem}"
        
        # Try to extract date from folder/filename
        folder_date = _extract_date_literal(combined_path_text)
        if folder_date:
            # If only year extracted (e.g., "2015"), convert to 01/01/YYYY
            if re.match(r'^\d{4}$', folder_date):
                folder_date = f"01/01/{folder_date}"
            
            # Use as effective_date if not already set
            if 'effective_date' not in metadata or metadata['effective_date'] == 'Unknown':
                metadata['effective_date'] = folder_date
                print(f"   📁 Effective Date from folder/filename: {folder_date}")
    
    # ENHANCEMENT: Use PDF creation date as fallback for effective_date if still Unknown
    if metadata.get('effective_date') == 'Unknown' and metadata_extra.get('pdf_creationdate'):
        pdf_date_str = metadata_extra['pdf_creationdate']
        # PDF dates format: D:20070806102619 → 2007-08-06
        if pdf_date_str.startswith('D:') and len(pdf_date_str) >= 10:
            year = pdf_date_str[2:6]
            month = pdf_date_str[6:8]
            day = pdf_date_str[8:10]
            metadata['effective_date'] = f"{month}/{day}/{year}"
            print(f"   📄 Effective Date from PDF metadata: {metadata['effective_date']}")

    # ENHANCEMENT: Company name and address extraction based on folder hierarchy
    # Structure: TopFolder (Company) / SubFolder (Agreement Type + Status) / Files
    path_obj = Path(file_path)
    
    # Get company name from TOP-LEVEL folder (grandparent or parent depending on depth)
    company_folder_name = None
    agreement_folder_name = None
    
    if len(path_obj.parts) >= 3:  # At least: root / company / agreement / file
        # Typical structure: Archive 2025 / Anixter, Inc / Reseller Agreement May 1998 - TERMINATED
        company_folder_name = path_obj.parts[-3]  # Grandparent
        agreement_folder_name = path_obj.parts[-2]  # Parent
    elif len(path_obj.parts) >= 2:
        # Simpler structure: Anixter, Inc / Reseller Agreement.pdf
        company_folder_name = path_obj.parts[-2]  # Parent is company
    
    # Extract counterparty (non-Spectralink company) info
    primary_party = metadata.get('primary_party_name') or metadata.get('company')
    counterparty_name = metadata.get('counterparty_name')
    
    # Check if counterparty looks like garbage (OCR noise)
    def _is_garbage_name(name: str) -> bool:
        if not name:
            return True
        # Check for OCR artifacts
        if re.search(r'#{2,}|`+|\\[A-Za-z]|<[A-Z]\\', name):
            return True
        # Check for excessive non-alpha characters
        alpha_ratio = sum(1 for c in name if c.isalpha()) / len(name) if name else 0
        if alpha_ratio < 0.5:
            return True
        # Check for file path patterns
        if re.search(r'\\[A-Za-z0-9_-]+\\', name):
            return True
        # Check for very long strings (likely OCR noise)
        if len(name) > 150:
            return True
        return False
    
    if _is_garbage_name(counterparty_name):
        counterparty_name = None
        metadata['counterparty_name'] = None
    if _is_garbage_name(primary_party):
        primary_party = None
        metadata['primary_party_name'] = None
    
    # Determine which party is Spectralink and which is counterparty
    is_primary_spectralink = primary_party and 'spectralink' in primary_party.lower()
    is_counter_spectralink = counterparty_name and 'spectralink' in counterparty_name.lower()
    
    # Check if this looks like a Spectralink contract (company folder is not Spectralink)
    # The folder structure tells us the counterparty: Archive 2025 / Anixter, Inc / Agreement
    is_company_folder_spectralink = company_folder_name and 'spectralink' in company_folder_name.lower()
    
    # If neither party detected but folder isn't Spectralink, assume Spectralink is primary
    if not is_primary_spectralink and not is_counter_spectralink and not is_company_folder_spectralink:
        # This is likely a Spectralink contract with the counterparty in the folder name
        is_primary_spectralink = True  # Assume Spectralink is primary
        primary_party = 'Spectralink Corporation'
        counterparty_name = company_folder_name
    
    if is_primary_spectralink or is_counter_spectralink:
        # This is a Spectralink contract
        spectralink_party = primary_party if is_primary_spectralink else counterparty_name
        other_party = counterparty_name if is_primary_spectralink else primary_party
        
        # If other_party was extracted but looks unreliable, prefer folder name
        if other_party and company_folder_name:
            # Check if folder name is better (cleaner, shorter, more reasonable)
            other_lower = other_party.lower()
            folder_lower = company_folder_name.lower()
            # Prefer folder if:
            # 1. Other party is very long (OCR might have grabbed too much)
            # 2. Other party contains numbers (likely OCR artifacts or dates)
            # 3. Other party contains special characters
            if (len(other_party) > 80 or
                re.search(r'\d{4}', other_party) or  # Contains year-like numbers
                'spectralink' not in folder_lower):  # Folder isn't Spectralink (confirms it's the counterparty)
                # But only if folder name looks like a company name
                if not folder_lower.startswith('archive') and len(company_folder_name) > 2:
                    other_party = company_folder_name
        
        # If no other party but we have folder name, use it
        if not other_party and company_folder_name:
            other_party = company_folder_name
        
        # Standardize Spectralink name
        spectralink_lower = spectralink_party.lower()
        if 'europe' in spectralink_lower or 'aps' in spectralink_lower:
            spectralink_standardized = 'Spectralink Europe ApS'
        else:
            spectralink_standardized = 'Spectralink Corporation'
        
        # Assign Spectralink address based on date
        effective_date = metadata.get('effective_date', 'Unknown')
        spectralink_addr = detect_spectralink_address(spectralink_party, effective_date, text)
        
        if is_primary_spectralink:
            metadata['primary_party_name'] = spectralink_standardized
            metadata['primary_party_address'] = spectralink_addr
            print(f"   🏢 Spectralink: {spectralink_standardized}")
            print(f"      Address: {spectralink_addr}")
            # Set counterparty - always use other_party which is already cleaned/prioritized
            if other_party:
                metadata['counterparty_name'] = other_party
                print(f"   🤝 Counterparty: {other_party}")
        else:
            metadata['counterparty_name'] = spectralink_standardized
            metadata['counterparty_address'] = spectralink_addr
            metadata['primary_party_name'] = company_folder_name or other_party
            print(f"   🏢 Company: {metadata['primary_party_name']}")
            print(f"   🤝 Counterparty: {spectralink_standardized} @ {spectralink_addr}")
        
        # Extract counterparty address from document text
        if other_party or company_folder_name:
            # Search for address near company name in text
            search_name = other_party or company_folder_name
            address_match = re.search(
                rf'{re.escape(search_name)}[^.{{}}]{'{0,200}'}?(\d+\s+[A-Z][A-Za-z\s,]+(?:Street|Ave|Avenue|Blvd|Road|Dr|Drive|Way|Lane)[^.{{}}]{{0,100}}?\d{{5}})',
                text,
                re.IGNORECASE | re.DOTALL
            )
            if address_match:
                counterparty_addr = address_match.group(1).strip()
                if is_primary_spectralink:
                    metadata['counterparty_address'] = counterparty_addr
                else:
                    metadata['primary_party_address'] = counterparty_addr
                print(f"      Address: {counterparty_addr}")
    else:
        # Non-Spectralink contract - use folder hierarchy
        if company_folder_name:
            # Clean folder name: remove "Archive 2025", dates, status keywords
            clean_company = company_folder_name
            if clean_company.lower().startswith('archive'):
                clean_company = re.sub(r'^archive\s*\d{4}\s*', '', clean_company, flags=re.IGNORECASE).strip()
            clean_company = re.sub(r'\b(terminated|draft|superseded)\b', '', clean_company, flags=re.IGNORECASE).strip()
            clean_company = re.sub(r'\s+', ' ', clean_company).strip(' -_,')
            
            if clean_company and len(clean_company) > 2:
                metadata['company_folder'] = clean_company
                metadata['primary_party_name'] = clean_company
                print(f"   🏢 Company (from folder): {clean_company}")
    
    # Extract agreement type and status from sub-folder name
    if agreement_folder_name:
        # Parse folder like "Reseller Agreement May 1998 - TERMINATED SEP 9 2016"
        # Check for status keywords
        folder_lower = agreement_folder_name.lower()
        if 'terminated' in folder_lower:
            metadata['contract_status'] = 'Terminated'
            metadata['is_terminated'] = True
            # Try to extract termination date
            term_date_match = re.search(r'terminated\s+([a-z]{3}\s+\d{1,2}\s+\d{4})', folder_lower, re.IGNORECASE)
            if term_date_match and metadata.get('termination_date') == 'Unknown':
                metadata['termination_date'] = term_date_match.group(1)
                print(f"   ✂️  Terminated: {metadata['termination_date']}")
        elif 'superseded' in folder_lower:
            metadata['contract_status'] = 'Archive'
            print(f"   📦 Status: Superseded")
        elif 'draft' in folder_lower:
            metadata['contract_status'] = 'Draft'
            metadata['is_draft'] = True
            print(f"   📝 Status: Draft")
        
        # Store agreement folder for reference
        metadata['agreement_folder'] = agreement_folder_name

    _ensure_standard_metadata(metadata)
    _ensure_boolean_metadata_defaults(metadata)

    # Add table summaries if present
    if tables:
        metadata["table_summary"] = f"{len(tables)} tables found"
        for i, table in enumerate(tables[:3]):  # First 3 tables
            if 'data' in table and table['data']:
                metadata[f"table_{i+1}_preview"] = str(table['data'][:2])  # First 2 rows
    
    metadata = _sanitize_metadata_dict(metadata)

    # Add to vector database
    success = vector_db.add_contract(contract_id, text, metadata)
    
    if success:
        # Update session history if available
        if tool_context:
            history = tool_context.state.get("analysis_history", [])
            history.append({
                "action": "upload",
                "contract_id": contract_id,
                "timestamp": datetime.now().isoformat()
            })
            tool_context.state["analysis_history"] = history[-10:]
        
        table_info = f"\n📊 Tables: {len(tables)} extracted" if tables else ""
        
        return f"""✅ Contract uploaded successfully!

Contract ID: {contract_id}
File: {Path(file_path).name} ({file_ext})
📝 Auto-detected Type: {detected_type}
👥 Auto-detected Parties: {detected_parties}
Size: {len(text):,} characters{table_info}

The contract is now searchable in the vector database."""
    else:
        return "❌ Failed to add contract to database"


def extract_contract_metadata(
    contract_text: str,
    tool_context: Optional[ToolContext] = None
) -> str:
    """
    Extract metadata with RAG enhancement.
    Searches for similar contracts to improve accuracy.
    """
    global vector_db
    
    # RAG: Search for similar contracts
    rag_context = ""
    if vector_db and EMBEDDINGS_AVAILABLE:
        similar = vector_db.search_similar_contracts(contract_text[:1000], n_results=3)
        if similar:
            rag_context = "\n\n🔍 SIMILAR CONTRACTS FOUND (for context):\n"
            for i, contract in enumerate(similar, 1):
                rag_context += f"{i}. {contract['metadata'].get('contract_type', 'Unknown')} - "
                rag_context += f"{contract['metadata'].get('parties', 'Unknown parties')}\n"
    
    # Extract parties
    parties_match = re.search(
        r'between\s+(.+?)\s+(?:and|&)\s+(.+?)(?:\.|,|\()',
        contract_text,
        re.IGNORECASE | re.DOTALL
    )
    if parties_match:
        party1, party2 = parties_match.groups()
    else:
        party1, party2 = "Party A", "Party B"
    
    # Extract dates
    date_pattern = r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
    dates = re.findall(date_pattern, contract_text)
    effective_date = dates[0] if dates else "Not found"
    
    # Extract defined terms
    defined_terms = extract_defined_terms(contract_text)
    
    # Extract renewal terms
    renewal_info = extract_renewal_terms(contract_text)
    
    # Update session state
    if tool_context:
        user_dept = tool_context.state.get("user_department", "Legal")
        history = tool_context.state.get("analysis_history", [])
        history.append({
            "action": "metadata_extraction",
            "timestamp": datetime.now().isoformat(),
            "department": user_dept
        })
        tool_context.state["analysis_history"] = history[-10:]
    
    return f"""📋 CONTRACT METADATA EXTRACTION

🏢 PARTIES:
   Party 1: {party1}
   Party 2: {party2}

📅 KEY DATES:
   Effective Date: {effective_date}

📝 DEFINED TERMS ({len(defined_terms)} found):
   {chr(10).join([f'   • {term["full_term"]} ("{term["short_term"]}")' for term in defined_terms[:5]])}

🔄 RENEWAL TERMS:
   Auto-Renewal: {'Yes' if renewal_info['has_auto_renewal'] else 'No'}
   Notice Period: {renewal_info['notice_period'] or 'Not specified'}
   Renewal Term: {renewal_info['renewal_term'] or 'Not specified'}
{rag_context}
"""


def analyze_contract_risks(
    contract_text: str,
    tool_context: Optional[ToolContext] = None
) -> str:
    """Analyze contract risks with RAG-enhanced context."""
    global vector_db
    
    # Get sensitivity level from session
    sensitivity = tool_context.state.get("sensitivity_level", "internal") if tool_context else "internal"
    
    # RAG: Find similar risky contracts
    rag_insights = ""
    if vector_db and EMBEDDINGS_AVAILABLE:
        similar = vector_db.search_similar_contracts(contract_text[:1000], n_results=3)
        if similar:
            rag_insights = "\n\n📊 RISK INSIGHTS FROM SIMILAR CONTRACTS:\n"
            for contract in similar:
                if 'risk_level' in contract['metadata']:
                    rag_insights += f"   • Similar {contract['metadata']['contract_type']}: "
                    rag_insights += f"{contract['metadata']['risk_level']} risk\n"
    
    risks = {
        "high": [],
        "medium": [],
        "low": []
    }
    
    # High-risk patterns
    if re.search(r"unlimited liability|no cap|uncapped", contract_text, re.IGNORECASE):
        risks["high"].append("Unlimited liability exposure")
    if re.search(r"automatically renew", contract_text, re.IGNORECASE):
        risks["high"].append("Auto-renewal without clear opt-out")
    
    # Medium-risk patterns
    if re.search(r"indemnif", contract_text, re.IGNORECASE):
        risks["medium"].append("Indemnification clauses present - review scope")
    if not re.search(r"limitation of liability", contract_text, re.IGNORECASE):
        risks["medium"].append("No explicit liability limitations found")
    
    # Low-risk patterns
    if re.search(r"termination for convenience", contract_text, re.IGNORECASE):
        risks["low"].append("Flexible termination options available")
    
    # Filter based on sensitivity level
    if sensitivity == "confidential":
        filtered_risks = risks  # Show all risks
    elif sensitivity == "internal":
        filtered_risks = {"high": risks["high"], "medium": risks["medium"], "low": []}
    else:  # public
        filtered_risks = {"high": risks["high"], "medium": [], "low": []}
    
    result = "⚠️  RISK ANALYSIS REPORT\n\n"
    
    for level in ["high", "medium", "low"]:
        if filtered_risks[level]:
            result += f"🔴 {level.upper()} RISK:\n"
            for risk in filtered_risks[level]:
                result += f"   • {risk}\n"
            result += "\n"
    
    result += rag_insights
    return result


def generate_contract_summary(
    contract_text: str,
    tool_context: Optional[ToolContext] = None
) -> str:
    """Generate department-specific summary with RAG context."""
    dept = tool_context.state.get("user_department", "Legal") if tool_context else "Legal"
    
    sections = find_contract_sections(contract_text)
    
    if dept == "Legal":
        return f"""⚖️  LEGAL DEPARTMENT SUMMARY

Key Legal Sections Identified:
{chr(10).join([f'   • {section.upper()}' for section in sections.keys()])}

Focus Areas for Legal Review:
   • Liability and indemnification clauses
   • Dispute resolution mechanisms
   • Governing law and jurisdiction
   • Confidentiality obligations
"""
    elif dept == "Sales":
        return f"""💼 SALES DEPARTMENT SUMMARY

Commercial Terms Overview:
   • Payment terms section: {'✅ Found' if 'payment_terms' in sections else '❌ Not found'}
   • Termination clauses: {'✅ Found' if 'term_and_termination' in sections else '❌ Not found'}

Key Considerations:
   • Revenue recognition timing
   • Performance milestones
   • Commission eligibility
"""
    else:  # Procurement
        return f"""📦 PROCUREMENT SUMMARY

Supplier Management:
   • Contract structure: {len(sections)} sections identified
   • Compliance review required: Yes

Focus Areas:
   • Delivery terms and SLAs
   • Quality standards
   • Pricing and payment schedules
"""


# ============================================================================
# SAFETY GUARDRAILS
# ============================================================================

def input_safety_guardrail(
    callback_context: CallbackContext,
    llm_request: LlmRequest
) -> Optional[LlmResponse]:
    """Before-model callback for input validation."""
    user_message = str(llm_request)
    
    # Check for PII
    pii_detected = detect_pii(user_message)
    if pii_detected:
        print(f"🛡️  PII detected in input: {', '.join(pii_detected)}")
        # In production, you might block or redact
    
    # Check for blocked keywords (if session context available)
    try:
        if hasattr(callback_context, 'session') and hasattr(callback_context.session, 'state'):
            blocked_keywords = callback_context.session.state.get("blocked_keywords", [])
        else:
            blocked_keywords = []
    except:
        blocked_keywords = []
    
    for keyword in blocked_keywords:
        if keyword.lower() in user_message.lower():
            print(f"🚫 Blocked keyword detected: {keyword}")
            return LlmResponse(text=f"Sorry, I cannot process requests containing '{keyword}'")
    
    # Check for prompt injection attempts
    injection_patterns = [
        r"ignore previous instructions",
        r"disregard.*rules",
        r"you are now"
    ]
    for pattern in injection_patterns:
        if re.search(pattern, user_message, re.IGNORECASE):
            print("🚨 Potential prompt injection detected")
            return LlmResponse(text="Invalid request detected. Please rephrase your query.")
    
    return None  # Proceed with request


def tool_safety_guardrail(
    tool: BaseTool,
    args: Dict[str, Any],
    tool_context: ToolContext
) -> Optional[str]:
    """Before-tool callback for tool execution validation."""
    
    # Check for restricted parties
    if "contract_text" in args:
        text = args["contract_text"]
        if "RESTRICTED_PARTY_NAME" in text:
            print("🚫 Restricted party detected in contract")
            return "Cannot process contracts involving restricted parties."
    
    # Enforce sensitivity level
    sensitivity = tool_context.state.get("sensitivity_level", "internal")
    if sensitivity == "public" and tool.name == "analyze_contract_risks":
        print("⚠️  Risk analysis limited for public sensitivity level")
        # Still allow but log the restriction
    
    return None  # Proceed with tool execution


# ============================================================================
# PARALLEL PROCESSING FOR BULK LOADING
# ============================================================================

async def process_contract_batch(
    file_paths: List[str],
    batch_size: int = 4,
    use_ocr: bool = False,
    base_folder: Optional[str] = None,
    force_reprocess: bool = False
) -> Dict[str, Any]:
    """Process multiple contracts with duplicate detection and detailed logging."""

    results: Dict[str, Any] = {
        "processed": 0,
        "failed": 0,
        "skipped": 0,
        "contract_ids": [],
        "ocr_used": 0,
    }

    base_path: Optional[Path] = None
    if base_folder:
        try:
            base_path = Path(base_folder).resolve()
        except Exception:
            base_path = None

    for i in range(0, len(file_paths), batch_size):
        batch = file_paths[i:i + batch_size]
        for file_path in batch:
            rel_label: Optional[str] = None
            if base_path:
                try:
                    rel_label = str(Path(file_path).resolve().relative_to(base_path))
                except ValueError:
                    rel_label = None

            try:
                status = upload_contract(
                    file_path,
                    use_ocr=use_ocr,
                    relative_label=rel_label,
                    force_reprocess=force_reprocess,
                )

                if status.startswith("✅"):
                    results["processed"] += 1
                    cid = _extract_contract_id_from_message(status)
                    if cid:
                        results["contract_ids"].append(cid)
                    if use_ocr:
                        results["ocr_used"] += 1
                elif "already ingested" in status:
                    print(status)
                    results["skipped"] += 1
                else:
                    print(status)
                    results["failed"] += 1
            except Exception as err:
                print(f"❌ Error processing {file_path}: {err}")
                results["failed"] += 1

        print(f"✅ Batch {i // batch_size + 1} complete: {len(batch)} files attempted")

    return results


# ============================================================================
# SEARCH TOOL FOR AGENT
# ============================================================================

def search_contracts_tool(
    query: str,
    n_results: int = 5,
    tool_context: Optional[ToolContext] = None
) -> str:
    """
    Search the contract database for relevant contracts.
    
    Args:
        query: Search query (e.g., "Keystone Logistics", "distributor agreements", "payment terms")
        n_results: Number of results to return (default 5)
    
    Returns:
        Formatted string with contract details and excerpts
    """
    global vector_db
    
    if not vector_db or not EMBEDDINGS_AVAILABLE:
        return "❌ Search unavailable - vector database not initialized"
    
    try:
        results = vector_db.search_similar_contracts(query, n_results=n_results)
        
        if not results:
            return f"No contracts found matching '{query}'"
        
        output = f"🔍 Found {len(results)} contracts matching '{query}':\n\n"
        
        for i, contract in enumerate(results, 1):
            hybrid_score = contract.get('hybrid_score', 0)
            metadata_score = contract.get('metadata_score', 0)
            semantic_score = contract.get('semantic_score', 0)
            
            output += f"{i}. {contract['id']}\n"
            output += f"   Match: {hybrid_score:.1f}% (Name/Party: {metadata_score:.0f}%, Content: {semantic_score:.1f}%)\n"
            output += f"   Type: {contract['metadata'].get('contract_type', 'Unknown')}\n"
            output += f"   Parties: {contract['metadata'].get('parties', 'Unknown')}\n"
            output += f"   Excerpt: {contract['text'][:300]}...\n\n"
        
        return output
    
    except Exception as e:
        return f"❌ Search error: {str(e)}"


# ============================================================================
# AGENT CREATION
# ============================================================================


def _select_llm_model():
    """Return the proper LLM implementation based on environment settings."""
    preferred = MODEL_PROVIDER
    wants_local = preferred in {"local", "ollama", "onprem", "offline"}

    if wants_local:
        if LOCAL_LLM_AVAILABLE:
            model_name = LOCAL_MODEL_NAME or "ollama"
            if LOCAL_MODEL_NAME:
                print(f"ℹ️  Using local Ollama profile '{LOCAL_MODEL_NAME}'.")
            else:
                print("ℹ️  Using local Ollama profile (config default).")
            return LocalOllamaLlm(model=model_name)
        else:
            print("⚠️  Local Ollama adapter unavailable; falling back to Gemini.")

    if preferred not in {"local", "ollama", "onprem", "offline"}:
        print(f"ℹ️  Using cloud Gemini model '{GEMINI_MODEL_NAME}'.")

    return Gemini(model=GEMINI_MODEL_NAME)


def create_clm_system():
    """Create the enhanced CLM agent with RAG capabilities."""
    if not ADK_AVAILABLE:
        raise RuntimeError("google-adk is required for agent/chat features.")
    
    llm_model = _select_llm_model()
    clm_manager = Agent(
        name="clm_rag_manager",
        model=llm_model,
        description="Enterprise CLM with ChromaDB RAG, PDF support, and parallel processing",
        instruction="""You are an Enterprise Contract Lifecycle Management assistant with RAG capabilities.

Your enhanced capabilities:
1. Upload PDF contracts and store them in ChromaDB vector database
2. Search contracts using semantic similarity
3. Extract metadata using semantic search for context
4. Analyze risks with insights from similar contracts
5. Generate summaries with department-specific focus

IMPORTANT: When a user asks about contracts (e.g., "summarize distributor agreements", "find Keystone contract"):
1. FIRST use search_contracts_tool to find relevant contracts
2. THEN analyze the returned contract text
3. Always search before trying to answer questions about specific contracts

When analyzing contracts, you automatically search for similar contracts to provide better context.

CRITICAL SECURITY GUIDELINES:
- Detect and flag PII in contracts
- Respect user sensitivity levels
- Log all operations for audit compliance
- Block restricted parties and keywords
""",
        tools=[
            search_contracts_tool,
            upload_contract,
            extract_contract_metadata,
            analyze_contract_risks,
            generate_contract_summary
        ],
        before_model_callback=input_safety_guardrail,
        before_tool_callback=tool_safety_guardrail,
        output_key="last_analysis_result"
    )
    
    return clm_manager


# ============================================================================
# MAIN EXECUTION
# ============================================================================

def main():
    global vector_db

    if not ADK_AVAILABLE:
        print("❌ google-adk not installed; interactive CLM console is unavailable.")
        print("   Install with: pip install google-adk")
        return
    
    print()
    print("=" * 70)
    print("🏢 ENTERPRISE CLM SYSTEM v2.0")
    print("=" * 70)
    print("📚 Features: ChromaDB Vector DB | Legal-BERT Embeddings | RAG-Powered AI")
    print("🚀 New: Main loop, AI Chat, Enhanced Search")
    print("=" * 70)
    print()
    
    # Check for pymupdf4llm (only show tip if NOT installed)
    if PYMUPDF4LLM_AVAILABLE:
        print("✅ pymupdf4llm: LLM-optimized PDF extraction ENABLED")
    else:
        print("💡 TIP: Install pymupdf4llm for greatly improved layout analysis:")
        print("   pip install pymupdf4llm")
    if RAPIDFUZZ_AVAILABLE:
        print("✅ rapidfuzz: C-accelerated fuzzy matching ENABLED")
    else:
        print("💡 TIP: Install rapidfuzz for higher-accuracy party detection:")
        print("   pip install rapidfuzz")
    print()
    
    # Check dependencies
    missing_deps = []
    if not CHROMADB_AVAILABLE:
        missing_deps.append("chromadb")
    if not PDF_SUPPORT:
        missing_deps.append("pymupdf")
    if not DOCX_SUPPORT:
        missing_deps.append("python-docx")
    if not EMBEDDINGS_AVAILABLE:
        missing_deps.append("sentence-transformers")
    if not OCR_SUPPORT:
        print("ℹ️  Optional: Install ocrmypdf for scanned PDF support")
    
    if missing_deps:
        print("⚠️  MISSING DEPENDENCIES:")
        print(f"   Install with: pip install {' '.join(missing_deps)}")
        print()
        print("   Without these, functionality will be limited:")
        print("   - No vector database storage (ChromaDB)")
        print("   - No PDF parsing (PyMuPDF)")
        print("   - No Word document support (python-docx)")
        print("   - No semantic search (sentence-transformers)")
        print()
        response = input("Continue anyway? (y/n): ")
        if response.lower() != 'y':
            return
    
    # Initialize ChromaDB (suppress Legal-BERT loading message)
    if CHROMADB_AVAILABLE:
        print("📁 Initializing ChromaDB vector database...")
        import warnings
        import sys
        from io import StringIO
        
        # Suppress warnings, stdout AND stderr during model loading
        warnings.filterwarnings("ignore")
        old_stdout = sys.stdout
        old_stderr = sys.stderr
        sys.stdout = StringIO()
        sys.stderr = StringIO()
        
        vector_db = ContractVectorDB()
        
        sys.stdout = old_stdout
        sys.stderr = old_stderr
        warnings.filterwarnings("default")
        
        print(f"✅ Embedding model: bert-base-uncased-contracts (trained specifically on contracts)")
        print(f"✅ Vector DB initialized: {vector_db.count_contracts()} contracts loaded")
    else:
        print("❌ Running without vector database")
    
    # Setup session service
    print("\n📁 Setting up session management...")
    session_service = InMemorySessionService()
    
    # Create session using asyncio.run() since create_session is async
    import asyncio
    session = asyncio.run(session_service.create_session(
        app_name="clm_rag_system",
        user_id="clm_user_001"
    ))
    
    # Initialize session state directly via .state dictionary
    session.state["user_department"] = "Legal"
    session.state["sensitivity_level"] = "internal"
    session.state["analysis_history"] = []
    session.state["blocked_keywords"] = ["BLOCK", "RESTRICTED_TERM"]
    session.state["user_preferences"] = {"temperature_unit": "celsius"}
    
    print("✅ Session created with default settings")
    print(f"   Session ID: {session.id}")
    
    # Create agent
    print("\n🏗️  Building CLM System with RAG...")
    clm_agent = create_clm_system()
    print("✅ CLM RAG Manager created!")
    
    # Helper function to create fresh runner (prevents event loop issues)
    def create_fresh_runner():
        return Runner(
            app_name="clm_rag_system",
            agent=clm_agent,
            session_service=session_service
        )
    
    # MAIN LOOP - Keep running until user exits
    while True:
        print("\n" + "=" * 70)
        print("Choose mode:")
        print("1. Upload PDF/DOCX contract")
        print("2. Upload with OCR (scanned PDFs)")
        print("3. Search contracts (semantic) + AI Analysis")
        print("4. Bulk upload folder")
        print("5. View database stats")
        print("6. Interactive AI Chat (ask questions about contracts)")
        print("0. Exit")
        print("=" * 70)
        
        mode = input("\nEnter 0-6: ").strip()
        
        if mode == "0":
            print("\n👋 Goodbye!")
            break
        
        elif mode == "1":
            file_path = input("Enter file path (PDF or DOCX): ").strip()
            
            result = upload_contract(file_path, use_ocr=False)
            print(f"\n{result}")
        
        elif mode == "2":
            if not OCR_SUPPORT:
                print("❌ OCR support requires ocrmypdf. Install with: pip install ocrmypdf")
                continue
            
            pdf_path = input("Enter scanned PDF path: ").strip()
            
            print("\n🔍 Processing with OCR (this may take 1-2 minutes)...")
            result = upload_contract(pdf_path, use_ocr=True)
            print(f"\n{result}")
    
        elif mode == "3":
            if not vector_db or not EMBEDDINGS_AVAILABLE:
                print("❌ Semantic search requires ChromaDB and sentence-transformers")
                continue
            
            query = input("Enter search query: ").strip()
            results = vector_db.search_similar_contracts(query, n_results=3)
            
            if not results:
                print("\n❌ No similar contracts found")
                continue
            
            print(f"\n🔍 Found {len(results)} similar contracts:\n")
            for i, contract in enumerate(results, 1):
                # Show hybrid scoring breakdown
                hybrid_score = contract.get('hybrid_score', 0)
                metadata_score = contract.get('metadata_score', 0)
                semantic_score = contract.get('semantic_score', 0)
                
                print(f"{i}. ID: {contract['id']}")
                print(f"   Type: {contract['metadata'].get('contract_type', 'Unknown')}")
                print(f"   Parties: {contract['metadata'].get('parties', 'Unknown')}")
                print(f"   Match: {hybrid_score:.1f}% (Name: {metadata_score:.0f}%, Content: {semantic_score:.1f}%)")
                print()
            
            # Ask if user wants AI analysis
            analyze = input("Run AI analysis on these contracts? (y/n): ").strip().lower()
            if analyze == 'y':
                print("\n🤖 Analyzing contracts with AI...\n")
                # Get the top contract's full text for analysis
                top_contract = results[0]
                contract_text = top_contract.get('text', '')[:5000]  # First 5000 chars
                
                # Create fresh session to prevent event loop issues
                analysis_session = asyncio.run(
                    session_service.create_session(
                        app_name="clm_rag_system",
                        user_id="clm_user_001"
                    )
                )
                analysis_session.state["user_department"] = session.state.get("user_department", "Legal")
                analysis_session.state["sensitivity_level"] = session.state.get("sensitivity_level", "internal")
                
                # Use fresh runner to get AI analysis with proper Content object
                analysis_prompt = f"Analyze this contract and answer: {query}\n\nContract excerpt:\n{contract_text}"
                message_content = Content(role="user", parts=[Part(text=analysis_prompt)])
                
                fresh_runner = create_fresh_runner()
                response_gen = fresh_runner.run(
                    session_id=analysis_session.id,
                    user_id=analysis_session.user_id,
                    new_message=message_content
                )
                
                # Iterate through generator and print response
                print("AI Analysis:\n")
                for event in response_gen:
                    if hasattr(event, 'content') and event.content:
                        for part in event.content.parts:
                            if hasattr(part, 'text') and part.text:
                                print(part.text, end="", flush=True)
                print("\n")
    
        elif mode == "4":
            folder_path_input = input("Enter folder path with contracts: ").strip()
            folder_path = Path(folder_path_input)
            pdf_files = list(folder_path.glob("*.pdf"))
            docx_files = list(folder_path.glob("*.docx")) + list(folder_path.glob("*.doc"))
            all_files = pdf_files + docx_files
            
            print(f"\nFound {len(pdf_files)} PDF files and {len(docx_files)} Word documents")
            use_ocr = input("Use OCR for scanned PDFs? (y/n): ").strip().lower() == 'y'
            force_reprocess = input("Reprocess duplicates? (y/N): ").strip().lower() == 'y'
            confirm = input("Process all? (y/n): ").strip()
            
            if confirm.lower() == 'y':
                print("\n🚀 Starting bulk upload...")
                results = asyncio.run(
                    process_contract_batch(
                        [str(f) for f in all_files],
                        use_ocr=use_ocr,
                        base_folder=str(folder_path.resolve()),
                        force_reprocess=force_reprocess,
                    )
                )
                print(f"\n✅ Bulk upload complete!")
                print(f"   Processed: {results['processed']}")
                print(f"   Skipped (already ingested): {results['skipped']}")
                print(f"   Failed: {results['failed']}")
                if use_ocr and results['ocr_used'] > 0:
                    print(f"   OCR applied: {results['ocr_used']} scanned documents")
        
        elif mode == "5":
            if not vector_db:
                print("❌ Vector database not available")
                continue
            
            count = vector_db.count_contracts()
            metadata = vector_db.get_all_metadata()
            
            print(f"\n📊 DATABASE STATISTICS")
            print(f"   Total contracts: {count}")
            print(f"\n📋 Recent uploads:")
            for item in metadata[-5:]:
                print(f"   • {item['contract_id']}: {item.get('contract_type', 'Unknown')}")
        
        elif mode == "6":
            print("\n💬 INTERACTIVE AI CHAT - Ask questions about your contracts!")
            print("💡 Examples:")
            print("   - 'Summarize all distributor agreements'")
            print("   - 'What are the payment terms in the Keystone contract?'")
            print("   - 'Find all contracts with renewal clauses'")
            print("   - 'Compare the NDAs we have'")
            print("\nCommands: /dept [Legal|Sales|Procurement], /sensitivity [public|internal|confidential]")
            print("Type 'exit' to return to main menu\n")
            
            # Create fresh chat session to prevent contamination from other modes
            chat_session = asyncio.run(
                session_service.create_session(
                    app_name="clm_rag_system",
                    user_id="clm_user_001"
                )
            )
            chat_session.state["user_department"] = session.state.get("user_department", "Legal")
            chat_session.state["sensitivity_level"] = session.state.get("sensitivity_level", "internal")
            
            while True:
                user_input = input("\n🗣️  You: ").strip()
                
                if user_input.lower() == 'exit':
                    break
                
                if user_input.startswith('/dept'):
                    dept = user_input.split()[1]
                    chat_session.state["user_department"] = dept
                    print(f"✅ Department set to: {dept}")
                    continue
                
                if user_input.startswith('/sensitivity'):
                    level = user_input.split()[1]
                    chat_session.state["sensitivity_level"] = level
                    print(f"✅ Sensitivity set to: {level}")
                    continue
                
                # Run through agent with AI (runner.run returns a generator)
                print("🤖 AI Agent: ", end="", flush=True)
                
                # Create proper Content object for message
                message_content = Content(role="user", parts=[Part(text=user_input)])
                
                # Use fresh runner to prevent event loop closure
                fresh_runner = create_fresh_runner()
                response_generator = fresh_runner.run(
                    session_id=chat_session.id,
                    user_id=chat_session.user_id,
                    new_message=message_content
                )
                
                # Iterate through the generator to get the actual response
                for event in response_generator:
                    if hasattr(event, 'content') and event.content:
                        for part in event.content.parts:
                            if hasattr(part, 'text') and part.text:
                                print(part.text, end="", flush=True)
                print()  # New line after response
        
        else:
            print("❌ Invalid mode selected")


if __name__ == "__main__":
    main()
